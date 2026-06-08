from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn

doc = Document()

# Page setup
for section in doc.sections:
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

style = doc.styles['Normal']
style.font.size = Pt(10)
style.font.name = 'Microsoft YaHei'
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
style.paragraph_format.line_spacing = 1.2

def add_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)
    p.paragraph_format.space_after = Pt(4)

def add_subtitle(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    p.paragraph_format.space_after = Pt(14)

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)
    return h

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    # Data rows
    for r, row_data in enumerate(rows):
        for c, val in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    doc.add_paragraph()  # spacer
    return table

# ========================================
# COVER
# ========================================
doc.add_paragraph()
add_title('RPA 项目管理文档模板')
add_subtitle('适用于企业内部 RPA 开发全流程 · 虞阅洲')
doc.add_paragraph()
doc.add_paragraph()

# ========================================
# 1. 需求调研表
# ========================================
add_heading('一、RPA 需求调研表', level=1)

p = doc.add_paragraph()
run = p.add_run('流程名称：BPM 月度工作量汇报自动填报')
run.font.size = Pt(10)
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('申请部门：综合管理部　　申请人/日期：张工 / 2025.03.10')
run.font.size = Pt(10)

add_heading('1. 流程概述', level=2)
doc.add_paragraph('当前操作：每月手动登录 BPM 系统，逐条录入 20+ 个项目的工作量数据。', style='List Bullet')
doc.add_paragraph('操作频率：每月 1 次，每次约 60 分钟。', style='List Bullet')
doc.add_paragraph('涉及系统：BPM 系统（Web 端）、月度项目清单 Excel。', style='List Bullet')

add_heading('2. 输入 / 输出', level=2)
doc.add_paragraph('输入数据：每月项目清单 Excel（项目名称、工时、进度百分比）。', style='List Bullet')
doc.add_paragraph('输出结果：BPM 系统填报完成，生成录入确认页面截图。', style='List Bullet')

add_heading('3. 异常分支', level=2)
doc.add_paragraph('部分项目可能已关闭，需自动跳过并标记原因。', style='List Bullet')
doc.add_paragraph('数据源 Excel 单元格为空时，弹出提示暂停执行。', style='List Bullet')

add_heading('4. 预期效果', level=2)
doc.add_paragraph('节省时间：约 50 分钟 / 月。', style='List Bullet')
doc.add_paragraph('准确率提升：从手动录入 ~95% → 自动化 99%+。', style='List Bullet')
doc.add_paragraph('减少人工：每月固定重复性操作不再占用人力。', style='List Bullet')

add_heading('5. 业务部门确认', level=2)
doc.add_paragraph('确认人签字：______________　　日期：______________')

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(4)
run = p.add_run('—— 以下由 IT 部门填写 ——')
run.italic = True
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

add_heading('6. IT 技术评估', level=2)
doc.add_paragraph('技术可行性：☑ 可行　　□ 部分可行　　□ 不可行', style='List Bullet')
doc.add_paragraph('预估工时：2 天（含测试与文档编写）。', style='List Bullet')
doc.add_paragraph('风险点：BPM 系统页面改版可能导致元素定位失效。', style='List Bullet')
doc.add_paragraph('评估人：______________　　日期：______________')

doc.add_page_break()

# ========================================
# 2. 方案审批表
# ========================================
add_heading('二、RPA 开发方案审批表', level=1)

add_table(
    ['项目', '内容'],
    [
        ['需求编号', 'RPA-2025-001'],
        ['流程名称', 'BPM 月度工作量汇报自动填报'],
        ['申请部门', '综合管理部'],
        ['开发方式', 'Python 3.10 + Selenium 4 + openpyxl'],
        ['异常处理', 'Step 级 try-except + 失败自动重试 3 次（间隔 5 秒）'],
        ['日志策略', 'Excel 日志表（日期 / 成功数 / 失败数 / 异常详情）'],
        ['通知方式', '报错自动邮件通知申请人 + IT 负责人'],
        ['预计工时', '2 个工作日'],
        ['预计上线', '2025 年 3 月 15 日'],
    ]
)

doc.add_paragraph()
doc.add_paragraph('审批意见：', style='List Bullet')
doc.add_paragraph()
doc.add_paragraph('□ 同意，按方案执行', style='List Bullet')
doc.add_paragraph('□ 修改后同意，修改意见：______________', style='List Bullet')
doc.add_paragraph('□ 不同意', style='List Bullet')
doc.add_paragraph()

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8)
run = p.add_run('IT 负责人签字：______________　　业务部门确认：______________')
run.font.size = Pt(9)

p = doc.add_paragraph()
run = p.add_run('日期：______________')
run.font.size = Pt(9)

doc.add_page_break()

# ========================================
# 3. 操作手册
# ========================================
add_heading('三、RPA 操作手册', level=1)

p = doc.add_paragraph()
run = p.add_run('流程名称：BPM 月度工作量汇报自动填报')
run.bold = True
p = doc.add_paragraph()
run = p.add_run('适用人员：综合管理部　　版本：v1.0')
p.paragraph_format.space_after = Pt(10)

add_heading('1. 使用步骤', level=2)

doc.add_paragraph('步骤一：打开桌面上的"数据源.xlsx"，按模板填写本月需要汇报的项目清单。', style='List Number')
doc.add_paragraph('步骤二：确认 BPM 系统可正常登录（浏览器自动操作，确保 Chrome 已安装）。', style='List Number')
doc.add_paragraph('步骤三：双击运行桌面的"工作量汇报.exe"，等待程序自动执行。', style='List Number')
doc.add_paragraph('步骤四：执行完成后弹出提示："完成！成功 18 条，失败 0 条"。', style='List Number')
doc.add_paragraph('步骤五：打开同目录下的"运行日志.xlsx"，核对结果无误即可。', style='List Number')

add_heading('2. 常见问题（FAQ）', level=2)

add_table(
    ['问题', '原因', '解决方法'],
    [
        ['提示"BPM 页面打开失败"', '网络异常或 VPN 未连接', '检查网络 → 确认 Chrome 能正常打开 BPM 登录页'],
        ['部分项目没有录入', '项目可能已关闭或被锁定', '查看日志表"跳过原因"列，确认是预期内的问题'],
        ['执行到一半卡住不动', 'BPM 页面弹出了意外弹窗', '截图发给 IT 处理；一般重启后可继续'],
        ['提示"Excel 文件未找到"', '数据源文件未放在指定位置', '将"数据源.xlsx"拷贝到脚本同目录下'],
    ]
)

add_heading('3. 注意事项', level=2)
doc.add_paragraph('运行期间请勿操作鼠标键盘，否则可能干扰模拟操作。', style='List Bullet')
doc.add_paragraph('建议在非高峰期运行（如午休时段），避免与其他操作冲突。', style='List Bullet')
doc.add_paragraph('每月 1 号上午为固定运行时间，如遇节假日顺延至下一个工作日。', style='List Bullet')

add_heading('4. 技术支持', level=2)

add_table(
    ['角色', '姓名', '联系方式'],
    [
        ['RPA 负责人', '虞阅洲', 'yuyuezhou1@gmail.com'],
        ['IT 部门', '（请补充）', '（请补充）'],
    ]
)

# Save
output = 'C:/Users/admin/Desktop/RPA项目管理文档模板_虞阅洲.docx'
doc.save(output)
print(f'Done: {output}')
