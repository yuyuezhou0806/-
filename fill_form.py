from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

doc = Document('C:/Users/admin/Desktop/01.应聘登记表（虞阅洲）.docx')

t1 = doc.tables[1]

def set_cell(cell, text):
    for p in cell.paragraphs:
        for run in p.runs:
            run.text = ''
    p = cell.paragraphs[0]
    if p.runs:
        p.runs[0].text = text
    else:
        run = p.add_run(text)
        run.font.size = Pt(9)

# Update Knowledge of Position (Table 1, Row 18, Cell 0)
set_cell(t1.rows[18].cells[0],
    'RPA开发工程师的核心是将企业重复性、规则化的工作流程自动化，释放人力去做更高价值的分析和决策工作。\n\n'
    '我在检测行业7年的综合管理工作中，深刻体会到物流/工程行业在单据处理、数据录入、跨系统对接等环节'
    '存在大量可被自动化的重复劳动。为此，我在现公司独立开发了BPM RPA自动化套件（18个脚本/3600行Python代码），'
    '用Selenium + Excel驱动方式覆盖了工作量汇报、商机登记、合同归档等核心流程，准确率>99%，每月节省15+小时。\n\n'
    '对于雅玛多国际物流的RPA岗位，我的优势在于：\n'
    '① 有成熟的RPA项目落地经验，从需求分析→脚本开发→异常处理→运维迭代全流程都能独立完成\n'
    '② 熟悉Python自动化技术栈（Selenium/openpyxl/pyautogui/win32com），能快速适配不同业务系统\n'
    '③ 有物流单据和流程管理经验（前公司涉及工艺规程、化学品台账等标准化流程），对物流行业的单据流转有实际认知\n'
    '④ 习惯用AI辅助开发（Claude Code等），开发效率是传统方式的3-5倍，能快速响应业务部门的自动化需求'
)

doc.save('C:/Users/admin/Desktop/01.应聘登记表（虞阅洲）.docx')
print('Updated!')
