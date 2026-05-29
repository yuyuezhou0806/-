#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""项目结算报审单批量填充工具"""

import os
import sys
import shutil
import argparse
from datetime import datetime
from docx import Document
import pandas as pd


def format_amount(value):
    if pd.isna(value):
        return ""
    v = float(value) * 10000
    if v == int(v):
        return f"{int(v)}"
    return f"{v:.2f}"


def fill_paragraph(paragraph, new_text):
    first_run = None
    for run in paragraph.runs:
        if first_run is None:
            first_run = run
        run.text = ""
    if first_run:
        first_run.text = new_text
    else:
        paragraph.add_run(new_text)


def fill_settlement_form(template_path, output_path, data):
    shutil.copy(template_path, output_path)
    doc = Document(output_path)

    total_workload = format_amount(data.get("实际发生总工作量(含税)", ""))
    test_fee = format_amount(data.get("实际应支付检测费（含税）", ""))
    confirmed_income = format_amount(data.get("已确认收入（含税）", ""))
    invoice_amount = format_amount(data.get("开票金额(含税)", ""))
    received_amount = format_amount(data.get("已收款金额（含税）", ""))

    def _get_str(key):
        val = data.get(key, "")
        if pd.isna(val):
            return ""
        return str(val).strip()

    payer = _get_str("付款单位")
    contract_no = _get_str("合同编号（结账号）")
    project_name = _get_str("项目名称")
    manager = _get_str("项目负责人")
    today = datetime.now().strftime("%Y.%m.%d")

    for para in doc.paragraphs:
        if para.text.strip().startswith("报审单编号"):
            fill_paragraph(para, "报审单编号：")

    table = doc.tables[0]
    cell = table.rows[0].cells[0]
    for pi, para in enumerate(cell.paragraphs):
        text = para.text
        if not text.strip():
            continue

        if text.strip().startswith("致："):
            new_text = f"致：{payer}" if payer else "致："
            fill_paragraph(para, new_text)

        elif "根据我单位与贵单位签订" in text:
            if "检测项目" in project_name:
                name_part = project_name
            elif project_name.endswith("检测"):
                name_part = project_name + "项目"
            elif project_name.endswith("项目") or project_name.endswith("工程"):
                name_part = project_name + "检测"
            else:
                name_part = project_name + "检测项目"
            new_text = (
                f"根据我单位与贵单位签订的{name_part}，"
                f"合同编号为：（{contract_no}），我单位已完成合同约定的全部工作量，"
                f"现报上结算报审单，请贵公司予以审核确认。"
            )
            fill_paragraph(para, new_text)

        elif "实际发生的总工作量" in text:
            new_text = f"本工程按合同约定单价计费，实际发生的总工作量为{total_workload}元。"
            fill_paragraph(para, new_text)

        elif "实际应支付检测费" in text:
            new_text = f"根据合同的规定，本工程实际应支付检测费为：{test_fee}元。"
            fill_paragraph(para, new_text)

        elif "确认收入" in text:
            new_text = f"（其中已确认收入：{confirmed_income}元"
            fill_paragraph(para, new_text)

        elif "开发票金额" in text:
            new_text = f"开发票金额：{invoice_amount}元  、已收款金额：{received_amount}元）"
            fill_paragraph(para, new_text)

        elif "项目负责人" in text:
            new_text = f"项目负责人：{manager}"
            fill_paragraph(para, new_text)

        elif "日期：" in text and pi >= 15:
            new_text = f"　　　　　　　　                                 日期：{today}                               "
            fill_paragraph(para, new_text)

    doc.save(output_path)


def main():
    parser = argparse.ArgumentParser(description="项目结算报审单批量填充")
    parser.add_argument("excel", nargs="?", help="Excel数据文件路径")
    parser.add_argument("--template", "-t", help="Word模板文件路径")
    parser.add_argument("--output", "-o", help="输出目录")
    args = parser.parse_args()

    script_dir = os.path.dirname(os.path.abspath(__file__))

    excel_path = args.excel
    if not excel_path:
        candidates = [
            os.path.join(script_dir, "报审单信息导入模版(1).xlsx"),
            os.path.join(script_dir, "data.xlsx"),
        ]
        for c in candidates:
            if os.path.exists(c):
                excel_path = c
                break
        if not excel_path:
            print("错误：未找到Excel文件")
            sys.exit(1)

    template_path = args.template
    if not template_path:
        candidates = [
            os.path.join(script_dir, "template.docx"),
            os.path.join(script_dir, "完工结算报审单.docx"),
        ]
        for c in candidates:
            if os.path.exists(c):
                template_path = c
                break
        if not template_path:
            print("错误：未找到Word模板")
            sys.exit(1)

    output_dir = args.output or os.path.join(script_dir, "output")
    os.makedirs(output_dir, exist_ok=True)

    if template_path.endswith(".doc"):
        docx_path = template_path.replace(".doc", ".docx")
        if not os.path.exists(docx_path):
            print("注意：检测到 .doc 模板，需要转换为 .docx...")
            try:
                import win32com.client
                word = win32com.client.Dispatch("Word.Application")
                word.Visible = False
                doc = word.Documents.Open(os.path.abspath(template_path))
                doc.SaveAs2(os.path.abspath(docx_path), FileFormat=16)
                doc.Close()
                word.Quit()
                template_path = docx_path
                print("转换完成")
            except Exception as e:
                print(f"转换失败：{e}")
                sys.exit(1)
        else:
            template_path = docx_path

    df = pd.read_excel(excel_path, engine="openpyxl")
    print(f"读取到 {len(df)} 条记录")

    for idx, row in df.iterrows():
        contract_no = str(row.get("合同编号（结账号）", "") or "").strip()
        project_name = str(row.get("项目名称", "") or "").strip()
        manager_val = str(row.get("项目负责人", "") or "").strip()
        safe_name = project_name[:20].replace("/", "-").replace("\\", "-")
        filename = f"{contract_no}_{safe_name}_{manager_val}.docx"
        output_path = os.path.join(output_dir, filename)
        fill_settlement_form(template_path, output_path, row.to_dict())
        print(f"  [OK] {filename}")

    print(f"\n全部完成，输出目录：{output_dir}")


if __name__ == "__main__":
    main()
