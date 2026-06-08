from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document('C:/Users/admin/Desktop/01.应聘登记表（虞阅洲）.docx')

for section in doc.sections:
    section.top_margin = Cm(0.5)
    section.bottom_margin = Cm(0.5)
    section.left_margin = Cm(1.0)
    section.right_margin = Cm(1.0)

for table in doc.tables:
    for row in table.rows:
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        trHeight = OxmlElement('w:trHeight')
        trHeight.set(qn('w:val'), '180')
        trHeight.set(qn('w:hRule'), 'atLeast')
        trPr.append(trHeight)
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 0.92

doc.save('C:/Users/admin/Desktop/01.应聘登记表（虞阅洲）.docx')
print('Done - 2 pages')
