"""生成《项目工作规范》RPA优化建议 Word文档"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

doc = Document()

# 全部宋体：正文 12pt，标题加粗，正文字号比标题小
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)
style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.font.bold = False

for lvl, size in [(1, 15), (2, 14), (3, 13)]:
    hs = doc.styles[f'Heading {lvl}']
    hs.font.name = '宋体'
    hs.font.size = Pt(size)
    hs.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    hs.font.bold = True

# === 标题 ===
title = doc.add_heading('《关于项目工作规范与违规管理条例》优化建议', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.name = '宋体'
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    run.font.bold = True

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run('—— 便捷办公 · 完整数据 · 减流程 ——')
run.font.name = '宋体'
run.font.size = Pt(12)
run.font.bold = False

doc.add_paragraph()

# === 总体判断 ===
doc.add_heading('一、总体判断', level=1)
p = doc.add_paragraph()
p.add_run('这份规范覆盖了项目全生命周期（立项→节点→工作量→收入→成本→资料→结算→催款），逻辑完整、权责清晰。').font.size = Pt(12)
p2 = doc.add_paragraph()
p2.add_run('但核心痛点是：').font.size = Pt(12)
items = [
    '流程推进靠"人盯"：每日自查、每周抽查、每月排查，大量时间花在"巡逻"而非"干活"',
    '数据一致性靠"人记"：四统一原则好，但靠人工核对四个系统，容易漏、容易错',
    '审批流转靠"人推"：结算单要人通知人领取，催款要人记住时间，链条一断全卡住',
    '资料管理靠"双份"：电子+纸质双归档，内部项目也照做，增加了不必要的工作量',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

# === 便捷办公 ===
doc.add_heading('二、便捷办公：减少重复操作', level=1)

suggestions_1 = [
    {
        'title': '1. "四统一"自动化校验（对应第五条）',
        'body': '规范要求"现场进度=系统数据=过程资料=结算数据"全程实时一致，目前依赖人工自查。\n'
                '建议：写一个自动化校验脚本，每天凌晨自动比对 ERP / 项目管理系统 / BPM 四个维度的数据，'
                '不一致的自动标红、生成差异清单推送给项目经理，不用人每天手动核对四套系统。'
    },
    {
        'title': '2. 结算流转去人工化（对应第十一条）',
        'body': '目前结算流转：项目经理编结算单→子公司检查→系统发起申请→用印→"通知"业务员→业务员"点领取"→3天内送甲方。\n'
                '建议：系统用印后自动通过微信/钉钉/邮件推送结算单 PDF 给业务员，附带电子回执。'
                '去掉"人工通知"和"人工点击领取"两步。超时未送达自动升级抄送主管。'
    },
    {
        'title': '3. 催款自动提醒（对应第十二条）',
        'body': '催款目前靠业务员主动关注合同付款节点和项目进度，容易遗漏。\n'
                '建议：系统接入合同付款节点数据，到期前 7 天自动推送提醒给业务员；'
                '逾期 15 天自动升级抄送部门负责人；逾期 30 天抄送财务总监。'
                '业务员只需记录催款结果，不用自己记催款时间。'
    },
    {
        'title': '4. 确认收入偏差实时看板（对应第八条）',
        'body': '规范要求"偏差超 5% 预警、超 10% 追责"，目前靠月报人工核对。\n'
                '建议：做一个简单的实时看板（Excel 或网页），三线对比"确认收入 / 预算收入 / 合同金额"，'
                '偏差超 5% 自动标黄、超 10% 自动标红并弹窗提醒项目经理。不需要等月报出来才知道出问题。'
    },
]

for s in suggestions_1:
    doc.add_heading(s['title'], level=2)
    doc.add_paragraph(s['body'])

# === 完整数据 ===
doc.add_heading('三、完整数据：减少人为遗漏', level=1)

suggestions_2 = [
    {
        'title': '5. 立项阶段合同量自动校验（对应第七条）',
        'body': '合同量核定目前靠项目经理人工对比销售合同，前期核算疏漏会引发后期收入争议。\n'
                '建议：写一个校验脚本，读取"合同清单 × 系统预算 × 历史同类项目均价"三方比对，'
                '偏差超 10% 自动标红提醒启动补量方案，从源头堵住缺口。'
    },
    {
        'title': '6. 资料完整性自检拦截（对应第十条）',
        'body': '资料要求"签字、盖章、日期、影像佐证"齐全，现在靠人工检查，经常漏了后面返工补。\n'
                '建议：系统上传资料时做前端校验——缺签字标红不让提交、缺日期弹窗提示、'
                '附件文件名与内容不符提醒改名。从入口拦截残缺资料，省得后续返工补件。'
    },
    {
        'title': '7. 成本费用到期自动对账（对应第九条）',
        'body': '劳务费、配合费等超过 30 天不予审批，但有时候是供应商发票晚到，不是项目组的错。\n'
                '建议：系统自动标记即将到期（25 天提醒）和已到期（30 天预警）的成本单据，'
                '同时留一个"延期申请"入口，说明原因即可延期至 60 天。'
                '一刀切改为缓冲机制，既不放松管控、也不误伤正常业务。'
    },
]

for s in suggestions_2:
    doc.add_heading(s['title'], level=2)
    doc.add_paragraph(s['body'])

# === 减流程 ===
doc.add_heading('四、减流程：去掉不必要的环节', level=1)

suggestions_3 = [
    {
        'title': '8. 电子+纸质双归档 → 分层归档（对应第十条第五款）',
        'body': '规范要求所有项目"电子档案+纸质档案双归档"，对内部项目而言是重复劳动。\n'
                '建议：分层归档——甲方明确要求纸质的才走双归档，内部管理类项目纯电子归档即可。'
                '一年省一箱纸、省一个档案室的空间，也省了打印装订理档的时间。'
    },
    {
        'title': '9. "每日自查"改成"异常推送"（对应第二十二条）',
        'body': '项目经理每日自查 + 业务人员每日自查，每天固定花 30 分钟做重复核对。\n'
                '建议：系统自动跑校验（见第 1 条），正常的不推送、异常的才推给对应人。'
                '把人从"巡逻流程"中解放出来，每天省出来的半小时可以用来推进实际工作。'
    },
    {
        'title': '10. 节点逾期流程简化（对应第六条第三款）',
        'body': '节点逾期后，责任人须提交调整变更说明原因、整改措施、完成时限。'
                '目前可能是口头沟通→填表→签字→上传的流程。\n'
                '建议：做成一键提交——系统自动拉取逾期节点信息，责任人只需在手机上选择逾期原因（下拉菜单）+ 预计完成日期，一键提交，自动流转审批。不用每次从头写说明。'
    },
]

for s in suggestions_3:
    doc.add_heading(s['title'], level=2)
    doc.add_paragraph(s['body'])

# === 实施路径 ===
doc.add_heading('五、实施路径建议', level=1)

doc.add_paragraph(
    '以上 10 条建议按实施难度和见效速度排序：'
)

steps = [
    ('第一阶段（1-2 周，零成本）：', '异常推送（第 9 条）、催款提醒（第 3 条）——写几个 Python 脚本 + Excel 看板就能跑，不碰系统。'),
    ('第二阶段（1 个月，低投入）：', '结算流转自动化（第 2 条）、资料上传校验（第 6 条）、确认收入看板（第 4 条）——需要前端加校验规则和简单 API。'),
    ('第三阶段（2-3 个月，需 IT 配合）：', '四统一自动校验（第 1 条）、合同量校验（第 5 条）、分层归档（第 8 条）——涉及多系统数据拉取和存储策略变更。'),
]

for phase, detail in steps:
    p = doc.add_paragraph()
    p.add_run(phase)
    p.add_run(detail)

# === 落款 ===
doc.add_paragraph()
doc.add_paragraph()
p_end = doc.add_paragraph()
p_end.add_run('以上建议供讨论稿参考，具体落地可根据集团 IT 资源和技术栈调整。')
p_end.add_run(f'\n\n2026 年 6 月 2 日')

# === 保存 ===
output = r'C:\Users\admin\Desktop\项目规范优化建议_RPA方向.docx'
doc.save(output)
print(f'Saved: {output}')
