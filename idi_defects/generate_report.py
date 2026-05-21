"""IDI 现场缺陷检查报告生成器

工作流:
  1. 用 search_defects.py 找到要检查的缺陷条目,记下它们的 global_id
  2. 用 --init 生成一份清单模板,把 defect_id / 轴线 / 照片路径填进去
  3. 用 --input <清单.json> --output <报告.docx> 生成现场报告

清单 JSON 字段说明:
  title       - 报告标题
  project     - 工程项目名
  site        - 地块/楼栋等
  inspector   - 检查人
  date        - 检查日期(字符串即可)
  items[]     - 检查到的缺陷列表
    defect_id - 在 defects.json 中的 global_id (如 D0001)
    location  - 现场部位/轴线描述(如 "1#楼 ① - ⑥ 轴交 ⓐ 轴 ")
    remark    - 现场补充说明(可空)
    photo     - 现场照片路径(可空)

生成的报告会按【风险期限/大类】自动分组,每条做成一个 5 列小表格,
版式参考源 IDI 清单(序号 / 现场照片 / 缺陷描述 / 纠正建议 / 部位备忘)。

示例:
  python generate_report.py --init items_template.json
  python generate_report.py --input items.json --output 检查报告.docx
"""

from __future__ import annotations

import json
import sys
from collections import OrderedDict
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

DATA_PATH = Path(__file__).parent / "data" / "defects.json"


# ---------- 数据加载 ----------

def load_defects_index() -> dict[str, dict]:
    if not DATA_PATH.exists():
        sys.exit(
            f"[!] 未找到 {DATA_PATH},请先运行: python extract_defects.py"
        )
    data = json.loads(DATA_PATH.read_text(encoding="utf-8"))
    return {d["global_id"]: d for d in data}


# ---------- 模板生成 ----------

def write_template(out_path: Path) -> None:
    today = date.today().isoformat()
    sample = {
        "title": "XX项目工程质量潜在缺陷检查报告",
        "project": "XX住宅项目",
        "site": "1#-3#楼",
        "inspector": "(检查人姓名)",
        "date": today,
        "items": [
            {
                "defect_id": "D0001",
                "location": "1#楼 ① - ⑥ 轴交 Ⓐ 轴 桩头",
                "remark": "现场实勘补充说明,可留空",
                "photo": "",
            },
            {
                "defect_id": "D0100",
                "location": "2#楼 屋面",
                "remark": "",
                "photo": "C:/path/to/photo.jpg",
            },
        ],
    }
    out_path.write_text(
        json.dumps(sample, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    print(f"已生成清单模板: {out_path}")
    print("修改 items 列表,把 defect_id 改成你需要的条目 ID(D0001…D0599),")
    print("填写 location/remark/photo 后,再运行 --input 生成报告。")


# ---------- 报告样式辅助 ----------

def set_cell_font(cell, *, bold=False, size=10, color=None):
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.name = "宋体"
            r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            r.font.size = Pt(size)
            r.bold = bold
            if color:
                r.font.color.rgb = RGBColor(*color)


def shade_cell(cell, rgb_hex: str) -> None:
    """给单元格设置背景色"""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        from docx.oxml import OxmlElement

        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), rgb_hex)


def set_cell_border(cell, **kwargs) -> None:
    """给单元格加全部边框"""
    from docx.oxml import OxmlElement

    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        el = tc_borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            tc_borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), "000000")


def add_paragraph(doc, text, *, bold=False, size=11, align=None, space_after=4):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    r.font.name = "宋体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    r.font.size = Pt(size)
    r.bold = bold
    p.paragraph_format.space_after = Pt(space_after)
    return p


# ---------- 报告生成 ----------

def build_report(
    items: list[dict],
    meta: dict,
    out_path: Path,
    index: dict[str, dict],
) -> None:
    doc = Document()

    # 页边距
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    # 标题
    add_paragraph(
        doc,
        meta.get("title", "工程质量潜在缺陷检查报告"),
        bold=True, size=18,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=12,
    )

    # 项目元信息行
    info_lines = [
        ("工程项目", meta.get("project", "")),
        ("地块/楼栋", meta.get("site", "")),
        ("检查日期", meta.get("date", "")),
        ("检查人", meta.get("inspector", "")),
        ("缺陷条目数", str(len(items))),
    ]
    info_tbl = doc.add_table(rows=len(info_lines), cols=2)
    info_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_tbl.autofit = False
    for ri, (k, v) in enumerate(info_lines):
        row = info_tbl.rows[ri]
        row.cells[0].text = k
        row.cells[1].text = v
        row.cells[0].width = Cm(3.5)
        row.cells[1].width = Cm(11.5)
        shade_cell(row.cells[0], "1F4E78")
        set_cell_font(row.cells[0], bold=True, color=(255, 255, 255))
        set_cell_font(row.cells[1])
        for cell in row.cells:
            set_cell_border(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    add_paragraph(doc, "", space_after=6)

    # 按 (period, major) 分组,保持出现顺序
    groups: "OrderedDict[tuple[str, str], list[dict]]" = OrderedDict()
    missing: list[str] = []
    for item in items:
        did = item.get("defect_id", "").strip()
        d = index.get(did)
        if not d:
            missing.append(did)
            continue
        key = (d["category_period"], d["category_major"])
        groups.setdefault(key, []).append((item, d))

    if missing:
        print(f"[!] 以下 defect_id 未在 defects.json 中找到,已跳过: {missing}")

    if not groups:
        sys.exit("[!] 清单中没有有效的条目,无法生成报告")

    # 分组渲染
    for (period, major), entries in groups.items():
        add_paragraph(
            doc,
            f"【{period}】{major}  (共 {len(entries)} 条)",
            bold=True, size=13,
            space_after=4,
        )
        for idx, (item, d) in enumerate(entries, start=1):
            render_defect_table(doc, idx, item, d)
            add_paragraph(doc, "", space_after=2)

    # 签字行
    add_paragraph(doc, "", space_after=12)
    sign_tbl = doc.add_table(rows=1, cols=2)
    sign_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    sign_tbl.cell(0, 0).text = f"检查人(签字): {meta.get('inspector', '')}"
    sign_tbl.cell(0, 1).text = f"日期: {meta.get('date', '')}"
    for c in sign_tbl.rows[0].cells:
        set_cell_font(c)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def render_defect_table(doc, idx: int, item: dict, defect: dict) -> None:
    """每条缺陷渲染为 4 行 2 列结构(子分类 / 描述 / 建议 / 现场图)。"""
    tbl = doc.add_table(rows=5, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False

    rows = [
        ("序号 / 子分项",
         f"#{idx}  {defect['category_minor']}  (原条目 {defect['global_id']})"),
        ("现场部位 / 轴线", item.get("location", "")),
        ("缺陷描述", defect["description"]),
        ("纠正与预防建议", defect["suggestion"]),
        ("现场补充", item.get("remark", "")),
    ]
    for ri, (k, v) in enumerate(rows):
        tbl.rows[ri].cells[0].text = k
        tbl.rows[ri].cells[1].text = v or ""
        tbl.rows[ri].cells[0].width = Cm(3.5)
        tbl.rows[ri].cells[1].width = Cm(13.5)
        shade_cell(tbl.rows[ri].cells[0], "D9E1F2")
        set_cell_font(tbl.rows[ri].cells[0], bold=True)
        set_cell_font(tbl.rows[ri].cells[1])
        for cell in tbl.rows[ri].cells:
            set_cell_border(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # 现场照片
    photo = (item.get("photo") or "").strip()
    if photo and Path(photo).exists():
        photo_row = tbl.add_row()
        photo_row.cells[0].text = "现场照片"
        shade_cell(photo_row.cells[0], "D9E1F2")
        set_cell_font(photo_row.cells[0], bold=True)
        set_cell_border(photo_row.cells[0])
        cell = photo_row.cells[1]
        set_cell_border(cell)
        try:
            cell.paragraphs[0].add_run().add_picture(photo, width=Cm(8))
        except Exception as e:
            cell.text = f"[图片插入失败: {e}]"
    elif photo:
        # 路径写了但文件不在
        photo_row = tbl.add_row()
        photo_row.cells[0].text = "现场照片"
        shade_cell(photo_row.cells[0], "D9E1F2")
        set_cell_font(photo_row.cells[0], bold=True)
        photo_row.cells[1].text = f"[未找到图片: {photo}]"
        for c in photo_row.cells:
            set_cell_border(c)
            set_cell_font(c)


# ---------- CLI ----------

def main() -> None:
    import argparse

    parser = argparse.ArgumentParser(
        description="生成现场质量缺陷检查报告 docx",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=__doc__,
    )
    parser.add_argument("--init", help="生成清单模板到指定路径")
    parser.add_argument("--input", help="输入清单 JSON")
    parser.add_argument(
        "--output", default="质量缺陷检查报告.docx", help="输出 docx 路径(默认在 cwd)"
    )
    args = parser.parse_args()

    if args.init:
        write_template(Path(args.init))
        return

    if not args.input:
        parser.error("必须指定 --input <清单.json> 或 --init <模板.json>")

    items_doc = json.loads(Path(args.input).read_text(encoding="utf-8"))
    items = items_doc.get("items", [])
    meta = {k: v for k, v in items_doc.items() if k != "items"}

    if not items:
        sys.exit("[!] items 为空,没有要写入的缺陷条目")

    index = load_defects_index()
    out_path = Path(args.output)
    build_report(items, meta, out_path, index)
    print(f"已生成报告: {out_path}  ({len(items)} 条缺陷条目)")


if __name__ == "__main__":
    main()
