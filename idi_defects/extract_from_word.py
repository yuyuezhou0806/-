import glob, os, json, re
from docx import Document
from docx.oxml.ns import qn
from PIL import Image

def extract_period(h1_text):
    """从H1标题提取期限"""
    if '十年' in h1_text:
        return '十年期'
    elif '五年' in h1_text:
        return '五年期'
    elif '两年' in h1_text:
        return '两年期'
    elif '机电' in h1_text:
        return '机电安装'
    else:
        return '十年期'

def extract_major(h2_text):
    """从H2标题提取大类"""
    # 去掉括号编号
    text = re.sub(r'[（(].*?[）)]', '', h2_text)
    return text.strip()

def get_cell_image_rids(cell):
    """获取单元格中的图片rId列表"""
    rids = []
    for para in cell.paragraphs:
        for run in para.runs:
            blips = run._element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip')
            for blip in blips:
                embed = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                if embed:
                    rids.append(embed)
    return rids

def main():
    path = r'C:\Users\admin\cc_test\idi_defects'
    files = glob.glob(os.path.join(path, '*.docx'))
    big_file = [f for f in files if os.path.getsize(f) > 100000000][0]

    doc = Document(big_file)

    # 1. 建立表格和标题的映射
    body = doc.element.body
    current_h1 = ''
    current_h2 = ''
    table_categories = []

    for child in body:
        tag = child.tag.split('}')[-1]
        if tag == 'p':
            pPr = child.find(qn('w:pPr'))
            if pPr is not None:
                pStyle = pPr.find(qn('w:pStyle'))
                if pStyle is not None:
                    style_val = pStyle.get(qn('w:val'))
                    text = ''.join(t.text for t in child.findall('.//' + qn('w:t')))
                    if style_val == '1':
                        current_h1 = text
                        current_h2 = ''
                    elif style_val == '2':
                        current_h2 = text
        elif tag == 'tbl':
            table_categories.append({
                'h1': current_h1,
                'h2': current_h2,
                'period': extract_period(current_h1),
                'major': extract_major(current_h2),
            })

    # 2. 获取所有图片part
    img_parts = {}
    for rel in doc.part.rels.values():
        if 'image' in rel.reltype:
            img_parts[rel.rId] = rel.target_part

    # 3. 准备输出目录
    imgs_dir = os.path.join(path, 'web', 'imgs_v2')
    os.makedirs(imgs_dir, exist_ok=True)
    for f in os.listdir(imgs_dir):
        os.remove(os.path.join(imgs_dir, f))

    # 4. 提取所有表格数据
    all_defects = []
    defect_id = 1
    img_counter = 1

    for table_idx, table in enumerate(doc.tables):
        if table_idx >= len(table_categories):
            break

        cat = table_categories[table_idx]
        rows = table.rows
        if len(rows) < 3:
            continue

        # 第0行=表头, 第1行=子类别
        subcategory = rows[1].cells[0].text.strip()

        for row_idx in range(2, len(rows)):
            row = rows[row_idx]
            cells = row.cells
            if len(cells) < 4:
                continue

            seq = cells[0].text.strip()
            if not seq:
                continue

            # 提取文字
            problem = cells[2].text.strip() if len(cells) > 2 else ''
            suggestion = cells[3].text.strip() if len(cells) > 3 else ''

            if not problem:
                continue

            # 提取图片（只取第一张）
            image_file = None
            for cell in cells:
                rids = get_cell_image_rids(cell)
                if rids:
                    rid = rids[0]
                    if rid in img_parts:
                        img_part = img_parts[rid]
                        content_type = img_part.content_type
                        if content_type.endswith('jpeg'):
                            ext = 'jpg'
                        else:
                            ext = content_type.split('/')[-1]
                        image_file = f'D{defect_id:04d}.{ext}'
                        img_path = os.path.join(imgs_dir, image_file)
                        with open(img_path, 'wb') as f:
                            f.write(img_part.blob)
                    break

            defect = {
                'category_period': cat['period'],
                'category_major': cat['major'],
                'category_minor': subcategory,
                'seq': seq,
                'problem': problem,
                'suggestion': suggestion,
                'global_id': f'D{defect_id:04d}',
            }
            if image_file:
                defect['image'] = image_file

            all_defects.append(defect)
            defect_id += 1

    # 5. 保存defects.json
    output_path = os.path.join(path, 'web', 'defects_v2.json')
    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(all_defects, f, ensure_ascii=False, indent=2)

    # 6. 统计
    print(f'Total defects: {len(all_defects)}')
    print(f'Total images: {len(os.listdir(imgs_dir))}')

    # 按期限统计
    period_counts = {}
    major_counts = {}
    for d in all_defects:
        p = d['category_period']
        m = d['category_major']
        period_counts[p] = period_counts.get(p, 0) + 1
        major_counts[m] = major_counts.get(m, 0) + 1

    print('\nBy period:')
    for p, c in sorted(period_counts.items()):
        print(f'  {p}: {c}')

    print('\nBy major (top 15):')
    for m, c in sorted(major_counts.items(), key=lambda x: -x[1])[:15]:
        print(f'  {m}: {c}')

if __name__ == '__main__':
    main()
