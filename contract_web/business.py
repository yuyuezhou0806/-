#!/usr/bin/env python3
"""
合同自动填充工具 (OCR版)
从网页截图自动识别项目信息，填入Word建设工程检测合同

用法:
    python contract_filler_ocr.py <截图图片路径>
    python contract_filler_ocr.py <截图图片路径> --template <合同模板路径>

示例:
    python contract_filler_ocr.py screenshot.png
    python contract_filler_ocr.py screenshot.png --template "空白合同.doc"
"""

import argparse
import json
import os
import re
import sys
import tempfile

# 关闭 paddle 的 oneDNN / PIR 加速,避免在部分 CPU 上触发
# "ConvertPirAttribute2RuntimeAttribute not support" 错误。
os.environ.setdefault('FLAGS_use_mkldnn', '0')
os.environ.setdefault('FLAGS_enable_pir_in_executor', '0')


def ocr_image(image_path):
    """使用PaddleOCR识别图片文字"""
    try:
        from paddleocr import PaddleOCR
    except ImportError:
        print("[X] PaddleOCR未安装，请先运行: pip install paddlepaddle paddleocr")
        sys.exit(1)

    print("[*] 正在识别图片文字，请稍候...")
    ocr = PaddleOCR(
        use_textline_orientation=True,
        lang='ch',
    )
    result = ocr.ocr(image_path)

    # 提取所有文字
    texts = []
    if result and result[0]:
        for line in result[0]:
            if line:
                text = line[1][0]  # 文字内容
                confidence = line[1][1]  # 置信度
                if confidence > 0.5:  # 过滤低置信度
                    texts.append(text)

    full_text = '\n'.join(texts)
    return full_text, texts


def extract_fields(text):
    """从OCR识别文本中提取关键字段"""
    fields = {}

    # 预处理1:把 OCR 经常把 label 和 value 拆成两行的情况合并。
    # 但要小心:如果"下一行"本身也是一个 label(以 : 结尾),不要合并,
    # 否则会把空 label 的位置串错。
    def _merge_label_value(s):
        out = []
        lines = s.split('\n')
        i = 0
        while i < len(lines):
            cur = lines[i]
            # 如果当前行以 : 或 : 结尾(label),且下一行存在且不是 label,合并
            if i + 1 < len(lines) and re.search(r'[：:]\s*$', cur):
                nxt = lines[i + 1].strip()
                if nxt and not re.search(r'[：:]\s*$', nxt):
                    out.append(cur.rstrip() + nxt)
                    i += 2
                    continue
            out.append(cur)
            i += 1
        return '\n'.join(out)

    # 预处理2:表格OCR把字段名和值拆成两行(无冒号),自动补上冒号
    def _merge_table_rows(s):
        known_labels = [
            '委托单位', '检测单位', '工程名称', '工程标段', '合同金额描述',
            '工程地址', '所属区县', '报建编号', '建设单位', '设计单位',
            '见证单位', '施工单位', '受监质监站', '检测项目信息',
            '建设地址', '建设地点', '监理单位', '勘察单位', '建筑面积',
            '合同价格', '总投资', '工程性质', '合同工期', '施工许可证号',
        ]
        out = []
        lines = s.split('\n')
        i = 0
        while i < len(lines):
            cur = lines[i].strip()
            if i + 1 < len(lines):
                nxt = lines[i + 1].strip()
                if cur in known_labels and nxt and nxt not in known_labels:
                    out.append(cur + '：' + nxt)
                    i += 2
                    continue
            out.append(cur)
            i += 1
        return '\n'.join(out)

    text = _merge_label_value(text)
    text = _merge_table_rows(text)

    # 报建编号
    m = re.search(r'报建编号[：:]\s*(\S+)', text)
    if m:
        fields['报建编号'] = m.group(1).strip()

    # 工程名称 / 项目名称(处理OCR跨行截断,如"地下室工\n程")
    for pattern in [r'工程名称[：:]\s*([^\n]+)', r'项目名称[：:]\s*([^\n]+)']:
        m = re.search(pattern, text)
        if m:
            v = m.group(1).strip()
            lines = text.split('\n')
            for i, line in enumerate(lines):
                if m.group(1) in line:
                    for j in range(i + 1, min(i + 3, len(lines))):
                        nxt = lines[j].strip()
                        # 续行条件:短(<=3字)、无冒号、不是常见label
                        if nxt and len(nxt) <= 3 and '：' not in nxt and ':' not in nxt \
                           and nxt not in ['建设单位', '工程名称', '建设地址', '建设规模',
                                           '合同工期', '合同价格', '参建单位']:
                            v += nxt
                        else:
                            break
                    break
            fields['工程名称'] = v
            break

    # 工程地址 / 建设地址 / 建设地点
    m = re.search(r'工程地址[：:]\s*([^\n]+)', text)
    if m:
        fields['建设地址'] = m.group(1).strip()
    else:
        for pattern in [r'建设地址[：:]\s*([^\n]+)', r'建设地点[：:]\s*([^\n]+)']:
            m = re.search(pattern, text)
            if m:
                fields['建设地址'] = m.group(1).strip()
                break

    # 建设单位(优先取"委托单位",这是合同甲方;次选"建设单位")
    for pattern in [r'委托单位[：:]\s*([^\n]+)', r'建设单位[：:]\s*([^\n]+)']:
        m = re.search(pattern, text)
        if m:
            v = m.group(1).strip()
            # 过滤掉"委托单位"后面跟着"检测单位"等相邻label的情况
            if '检测单位' not in v and '设计单位' not in v and '施工单位' not in v:
                fields['建设单位'] = v
                break

    # 清理表格OCR跨列噪音的辅助函数
    def _clean_table_value(v):
        if not v:
            return v
        # 施工许可证表格中，字段值后面常跟着"项目负责人""总监理工程师"等，需要截断
        noise_keywords = ['项目负责人', '总监理工程师', '项目总监', '专业负责人']
        for noise in noise_keywords:
            if noise in v:
                v = v[:v.index(noise)].strip()
                break
        # 如果值太长且包含其他单位label，说明OCR跨行了
        if len(v) > 25:
            unit_keywords = ['勘察单位', '设计单位', '施工单位', '监理单位',
                             '检测单位', '总承包单位', '建设单位', '见证单位']
            for kw in unit_keywords:
                idx = v.find(kw)
                if idx > 0:  # 不是开头出现的（即后面混入了其他单位）
                    v = v[:idx].strip()
                    break
        return v

    # 表格字段提取辅助:支持OCR乱序(表格OCR常把label和值的顺序打乱)
    def _extract_table_field(label, text):
        lines = text.split('\n')
        # 先转义label中的正则特殊字符
        label_escaped = re.escape(label)

        for i, line in enumerate(lines):
            if label not in line:
                continue

            # 格式1: 同一行有label和值 (label:值)
            m = re.search(rf'{label_escaped}[：:]\s*([^\n]+)', line)
            if m:
                v = _clean_table_value(m.group(1).strip())
                if v and len(v) > 2:
                    return v

            # 格式2: label在后,值在前几行(OCR乱序常见于表格)
            # 向前查找最近的有效值,跳过噪音行
            candidates = []
            for j in range(i - 1, max(i - 6, -1), -1):
                v = lines[j].strip()
                if not v or v in ['项目负责人', '总监理工程师', '项目经理', '']:
                    continue
                # 判断是否是公司名/机构名(含特定关键词)
                if any(kw in v for kw in ['公司', '研究院', '集团', '事务所', '中心', '大学']):
                    candidates.append((j, v))

            if candidates:
                j, v = candidates[0]  # 取最近的一个
                # 处理公司名被拆行的情况,如"有"+"限公司"
                if v.endswith('有') and j + 1 < len(lines) and '限公司' in lines[j + 1]:
                    v += '限公司'
                # 处理"XX(集团)有"+"限公司"
                elif v.endswith(')') and '有' in v[-3:] and j + 1 < len(lines) and '限公司' in lines[j + 1]:
                    v += '限公司'
                return _clean_table_value(v)

        return None

    # 工程总承包单位(施工许可证中可能独立列出)
    for label in ['工程总承包单位', '总承包单位']:
        v = _extract_table_field(label, text)
        if v:
            fields['总承包单位'] = v
            break

    # 施工单位
    v = _extract_table_field('施工单位', text)
    if v:
        fields['施工单位'] = v

    # 设计单位
    v = _extract_table_field('设计单位', text)
    if v:
        fields['设计单位'] = v

    # 监理单位
    v = _extract_table_field('监理单位', text)
    if v:
        fields['监理单位'] = v

    # 见证单位
    v = _extract_table_field('见证单位', text)
    if v:
        fields['见证单位'] = v

    # 勘察单位
    v = _extract_table_field('勘察单位', text)
    if v:
        fields['勘察单位'] = v

    # 检测单位(签合同的乙方)
    v = _extract_table_field('检测单位', text)
    if v:
        fields['检测单位'] = v

    # 工程性质(注意:OCR 经常漏识别 form widget 里的"修缮""新建"等值,
    # 此时这个字段会是空 / 是后面别的 label。所以只取非空且不含 ":" 的值。
    m = re.search(r'工程性质[：:]\s*([^\n]+)', text)
    if m:
        v = m.group(1).strip()
        if v and '：' not in v and ':' not in v:
            fields['工程性质'] = v

    # 建筑面积 - 多种格式(报送系统/施工许可证)
    for pattern in [
        r'建筑面积[：:]\s*([\d.]+)',
        r'建设规模.*建筑面积[：:]\s*([\d.]+)',
        r'建筑面积\D*?([\d.]+)',
        r'总建筑面积\D*?([\d.]+)',
        r'建设规模.*总建筑面积\D*?([\d.]+)',
    ]:
        m = re.search(pattern, text)
        if m:
            fields['建筑面积'] = m.group(1).strip()
            break

    # 合同价格 / 总投资
    for pattern in [
        r'合同价格\(万元\)[：:]\s*([\d.]+)',
        r'总投资\(万元\)[：:]\s*([\d.]+)',
        r'合同价格[：:]\s*([\d.]+)',
        r'总投资[：:]\s*([\d.]+)',
    ]:
        m = re.search(pattern, text)
        if m:
            fields['合同价格'] = m.group(1).strip()
            break

    # 所属区县 - 优先用 explicit "工程所属区县:" label;再 fall back 用工程地址
    districts = ['奉贤区', '浦东新区', '徐汇区', '黄浦区', '静安区', '长宁区',
                 '普陀区', '虹口区', '杨浦区', '闵行区', '宝山区', '嘉定区',
                 '金山区', '松江区', '青浦区', '崇明区']

    m = re.search(r'工程所属区县[：:]\s*([^\n]+)', text)
    if m:
        v = m.group(1).strip()
        # 确保是合法区名
        for district in districts:
            if district in v or district.rstrip('区') in v:
                fields['所属区县'] = district
                break

    if '所属区县' not in fields:
        # fall back: 在工程地址里找区(只查地址字段,不查全文)
        addr = fields.get('建设地址', '') + fields.get('工程名称', '')
        for district in districts:
            if district in addr:
                fields['所属区县'] = district
                break

    # 质监站/监管机构 - 从发证机关推导
    m = re.search(r'发证机关[：:]\s*([^\n]+)', text)
    if m:
        org = m.group(1).strip()
        if '建设' in org and '管理' in org:
            fields['质监站'] = org.replace('建设和管理委员会', '建设工程安全质量监督站')
        else:
            fields['质监站'] = org

    # 施工许可证号(多种格式)
    for pattern in [
        r'施工许可证号[：:]\s*(\w+)',
        r'施工许可证编号[：:]\s*(\w+)',
        r'许可证编号[：:]\s*(\w+)',
        r'编号[：:]\s*(\w+)',
    ]:
        m = re.search(pattern, text)
        if m:
            fields['施工许可证号'] = m.group(1).strip()
            break

    # 合同工期
    m = re.search(r'合同工期[：:]\s*([^\n]+)', text)
    if m:
        fields['合同工期'] = m.group(1).strip()

    # 工程标段
    m = re.search(r'工程标段[：:]\s*([^\n]+)', text)
    if m:
        fields['工程标段'] = m.group(1).strip()

    # 合同金额描述
    m = re.search(r'合同金额描述[：:]\s*([^\n]+)', text)
    if m:
        v = m.group(1).strip()
        if v and v != '0':
            fields['合同金额描述'] = v

    # 受监质监站（直接从字段提取，不从发证机关推导）
    m = re.search(r'受监质监站[：:]\s*([^\n]+)', text)
    if m:
        fields['质监站'] = m.group(1).strip()

    # 检测项目信息
    m = re.search(r'检测项目信息[：:]\s*([^\n]+)', text)
    if m:
        fields['检测项目信息'] = m.group(1).strip()

    return fields


def build_contract_config(fields, template_path, output_name, check_type=""):
    """从提取的字段构建合同填充配置"""
    config = {
        'template_path': template_path,
        'output_name': output_name,
        'fields': {},
        'global_replaces': {},
    }

    # 检测模板中label的实际格式(带冒号还是不带),避免重复替换
    from docx import Document
    _doc = Document(template_path)
    _all_text = '\n'.join(p.text for p in _doc.paragraphs)
    for table in _doc.tables:
        for row in table.rows:
            for cell in row.cells:
                _all_text += '\n' + cell.text

    def _label(base):
        """返回模板中实际使用的label格式"""
        if base + '：' in _all_text or base + ':' in _all_text:
            return base + '：'
        if base in _all_text:
            return base
        return base + '：'  # 默认

    # 建设单位 -> 甲方、建设/实施单位、见证单位
    if '建设单位' in fields:
        company = fields['建设单位']
        config['fields']['甲方（委托单位）'] = f'甲方（委托单位） {company}'
        config['fields']['建设/实施单位：'] = f'建设/实施单位：{company}'
        config['fields']['见证单位：'] = f'见证单位：{company}'

    # 施工单位(根据模板格式设置key,避免重复替换)
    if '施工单位' in fields:
        label = _label('施工单位')
        config['fields'][label] = f'{label}{fields["施工单位"]}'

    # 总承包单位(施工许可证中可能独立列出,优先使用)
    if '总承包单位' in fields:
        label = _label('总承包单位')
        config['fields'][label] = f'{label}{fields["总承包单位"]}'
    elif '施工单位' in fields:
        label = _label('总承包单位')
        config['fields'][label] = f'{label}{fields["施工单位"]}'

    # 设计单位
    if '设计单位' in fields:
        label = _label('设计单位')
        config['fields'][label] = f'{label}{fields["设计单位"]}'

    # 监理单位
    if '监理单位' in fields:
        label = _label('监理单位')
        config['fields'][label] = f'{label}{fields["监理单位"]}'

    # 勘察单位
    if '勘察单位' in fields:
        label = _label('勘察单位')
        config['fields'][label] = f'{label}{fields["勘察单位"]}'

    # 见证单位
    if '见证单位' in fields:
        config['fields']['见证单位：'] = f'见证单位：{fields["见证单位"]}'

    # 工程名称
    if '工程名称' in fields:
        config['fields']['工程名称：'] = f'工程名称：{fields["工程名称"]}'

    # 工程地址
    if '建设地址' in fields:
        config['fields']['工程地址：'] = f'工程地址：{fields["建设地址"]}'

    # 报建编号（施工许可证中没有报建编号，用施工许可证号代替）
    if '报建编号' in fields:
        config['fields']['工程报建编号：'] = f'工程报建编号：{fields["报建编号"]}'
    elif '施工许可证号' in fields:
        config['fields']['工程报建编号：'] = f'工程报建编号：{fields["施工许可证号"]}'

    # 所属区县
    if '所属区县' in fields:
        config['fields']['工程所属区县：'] = f'工程所属区县：{fields["所属区县"]}'

    # 建筑面积
    if '建筑面积' in fields:
        config['fields']['建筑面积（㎡）：'] = f'建筑面积（㎡）：{fields["建筑面积"]}'

    # 投资额/建安费
    if '合同价格' in fields:
        price = fields['合同价格'] + '万元'
        config['fields']['工程投资额：'] = f'工程投资额：{price}'
        config['fields']['工程建安费：'] = f'工程建安费：{price}'

    # 质监站
    if '质监站' in fields:
        config['fields']['质监站/监管机构：'] = f'质监站/监管机构：{fields["质监站"]}'

    # 检测单位（乙方）
    if '检测单位' in fields:
        config['fields']['检测单位：'] = f'检测单位：{fields["检测单位"]}'
        config['fields']['乙方：'] = f'乙方：{fields["检测单位"]}'

    # 工程标段
    if '工程标段' in fields:
        config['fields']['工程标段：'] = f'工程标段：{fields["工程标段"]}'

    # 合同金额描述
    if '合同金额描述' in fields:
        config['fields']['合同金额描述：'] = f'合同金额描述：{fields["合同金额描述"]}'

    # 检测项目信息
    if '检测项目信息' in fields:
        config['fields']['检测项目信息：'] = f'检测项目信息：{fields["检测项目信息"]}'

    # 检测类别复选框处理
    # Word 模板里的复选框: ☐(未选中) ☑(已选中)
    # 先全部重置为未选中,再把用户选择的设为选中
    if check_type:
        selected = [t.strip() for t in check_type.split(',') if t.strip()]
        all_types = ['验收检测', '平行检测', '其他']
        for t in all_types:
            # 重置: 已选中 → 未选中
            config['global_replaces'][f'☑ {t}'] = f'☐ {t}'
        for t in selected:
            if t in all_types:
                # 选中: 未选中 → 已选中
                config['global_replaces'][f'☐ {t}'] = f'☑ {t}'

    return config


def fill_contract(config):
    """使用Word COM接口填入合同"""
    template_path = config.get('template_path', '')
    output_name = config.get('output_name', '合同-已填充.doc')
    fields = config.get('fields', {})
    global_replaces = config.get('global_replaces', {})

    if not os.path.exists(template_path):
        print(f"[X] 模板文件不存在: {template_path}")
        return False

    template_dir = os.path.dirname(template_path) or '.'
    output_path = os.path.join(template_dir, output_name)

    from docx import Document as DocxDocument

    def replace_in_para(para, old, new):
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        text = para.text
        if old not in text:
            return False

        # 检测段落中是否有下划线格式(任意run)
        para_has_underline = any(r.font.underline for r in para.runs)

        # 情况1: old 完全在单个 run 内
        for run in para.runs:
            if old in run.text:
                run_text = run.text
                pos = run_text.index(old)
                prefix = run_text[:pos]
                suffix = run_text[pos + len(old):]

                # 将 new 拆分为 label 部分和值部分
                # new 格式: "label 值" 或 "label:值" 等
                label_part = old  # label 就是 old 本身
                value_part = new[len(old):]  # new 中 old 之后的部分

                # 清空原 run
                run.text = ''

                run_elem = run._element
                last_elem = run_elem

                # 添加 prefix（保留原格式）
                if prefix:
                    r = OxmlElement('w:r')
                    t = OxmlElement('w:t')
                    t.text = prefix
                    r.append(t)
                    last_elem.addnext(r)
                    last_elem = r

                # 添加 label run（无下划线）
                if label_part:
                    r = OxmlElement('w:r')
                    t = OxmlElement('w:t')
                    t.text = label_part
                    r.append(t)
                    last_elem.addnext(r)
                    last_elem = r

                # 添加值 run（如果段落有下划线则加下划线）
                if value_part:
                    r = OxmlElement('w:r')
                    if para_has_underline:
                        rPr = OxmlElement('w:rPr')
                        u = OxmlElement('w:u')
                        u.set(qn('w:val'), 'single')
                        rPr.append(u)
                        r.append(rPr)
                    t = OxmlElement('w:t')
                    t.text = value_part
                    r.append(t)
                    last_elem.addnext(r)
                    last_elem = r

                # 添加 suffix（保留原格式）
                if suffix:
                    r = OxmlElement('w:r')
                    if run.font.underline:
                        rPr = OxmlElement('w:rPr')
                        u = OxmlElement('w:u')
                        u.set(qn('w:val'), 'single')
                        rPr.append(u)
                        r.append(rPr)
                    t = OxmlElement('w:t')
                    t.text = suffix
                    r.append(t)
                    last_elem.addnext(r)

                return True

        # 情况2: old 跨越多个 runs
        parts = [r.text for r in para.runs]
        full = ''.join(parts)
        pos = full.index(old)
        end_pos = pos + len(old)

        cumulative = 0
        start_idx = end_idx = -1
        for i, part in enumerate(parts):
            part_start = cumulative
            part_end = cumulative + len(part)
            if start_idx == -1 and part_start <= pos < part_end:
                start_idx = i
            if start_idx != -1 and part_end >= end_pos:
                end_idx = i
                break
            cumulative = part_end

        if start_idx == -1 or end_idx == -1:
            return False

        prefix_len = pos - sum(len(parts[j]) for j in range(start_idx))
        prefix = parts[start_idx][:prefix_len] if prefix_len > 0 else ''
        suffix_start = end_pos - sum(len(parts[j]) for j in range(end_idx))
        suffix = parts[end_idx][suffix_start:] if suffix_start < len(parts[end_idx]) else ''

        # 分离 label 和值
        label_part = old
        value_part = new[len(old):]

        # 清空涉及的runs(保留下划线runs)
        for i in range(start_idx, end_idx + 1):
            run_text = para.runs[i].text
            is_underline_run = (para.runs[i].font.underline and
                                len(run_text) > 0 and
                                all(c in ' \t' for c in run_text))
            if not is_underline_run:
                para.runs[i].text = ''

        # 在 start run 写入 prefix + label（无下划线）
        para.runs[start_idx].text = prefix + label_part
        para.runs[start_idx].font.underline = False

        last_elem = para.runs[start_idx]._element

        # 添加值 run（带下划线）
        if value_part:
            r = OxmlElement('w:r')
            if para_has_underline:
                rPr = OxmlElement('w:rPr')
                u = OxmlElement('w:u')
                u.set(qn('w:val'), 'single')
                rPr.append(u)
                r.append(rPr)
            t = OxmlElement('w:t')
            t.text = value_part
            r.append(t)
            last_elem.addnext(r)
            last_elem = r

        # 添加 suffix run（有下划线）
        if suffix:
            r = OxmlElement('w:r')
            if any(para.runs[i].font.underline for i in range(start_idx, end_idx + 1)):
                rPr = OxmlElement('w:rPr')
                u = OxmlElement('w:u')
                u.set(qn('w:val'), 'single')
                rPr.append(u)
                r.append(rPr)
            t = OxmlElement('w:t')
            t.text = suffix
            r.append(t)
            last_elem.addnext(r)

        return True

    try:
        doc = DocxDocument(template_path)
        print(f"[*] 已打开模板: {os.path.basename(template_path)}")

        success = 0
        fail = 0
        for field_label, value in fields.items():
            found = False
            for para in doc.paragraphs:
                if replace_in_para(para, field_label, value):
                    found = True
                    break
            if not found:
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                if replace_in_para(para, field_label, value):
                                    found = True
                                    break
                            if found: break
                        if found: break
                    if found: break
            if found:
                success += 1
                print(f"  [OK] {field_label} -> {value[:40]}")
            else:
                fail += 1
                print(f"  [X] 未找到: {field_label}")

        for old, new in global_replaces.items():
            for para in doc.paragraphs:
                replace_in_para(para, old, new)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            replace_in_para(para, old, new)

        doc.save(output_path)
        print(f"\n[OK] 合同已保存: {output_path}")
        print(f"     成功填入 {success} 个字段")
        if fail > 0:
            print(f"     {fail} 个字段未找到（模板中可能不存在）")
        return True

    except Exception as e:
        print(f"[X] 错误: {e}")
        import traceback
        traceback.print_exc()
        return False


def main():
    parser = argparse.ArgumentParser(description='从网页截图自动识别并填入建设工程检测合同')
    parser.add_argument('image', help='网页截图图片路径')
    parser.add_argument('--template', '-t', default=r'C:\Users\admin\Desktop\合同生成\空白合同【房建】（2021版）.doc',
                        help='合同模板路径')
    parser.add_argument('--output', '-o', help='输出文件名（默认从工程名称生成）')
    parser.add_argument('--save-config', '-s', action='store_true',
                        help='保存生成的config.json到临时目录')
    parser.add_argument('--dry-run', '-d', action='store_true',
                        help='只显示识别结果，不填入合同')

    args = parser.parse_args()

    if not os.path.exists(args.image):
        print(f"[X] 图片文件不存在: {args.image}")
        sys.exit(1)

    # Step 1: OCR识别
    full_text, texts = ocr_image(args.image)

    print("\n" + "=" * 60)
    print("  OCR识别结果")
    print("=" * 60)
    for i, text in enumerate(texts, 1):
        print(f"  {i}. {text}")

    # Step 2: 提取字段
    print("\n" + "=" * 60)
    print("  提取的字段")
    print("=" * 60)
    fields = extract_fields(full_text)
    if not fields:
        print("  [!] 未能从图片中提取到字段")
        sys.exit(1)

    for key, value in fields.items():
        print(f"  {key}: {value}")

    # Step 3: 构建配置
    project_name = fields.get('工程名称', '合同')
    safe_name = re.sub(r'[\\/:*?"<>|]', '_', project_name)
    output_name = args.output or f"{safe_name}-已填充.doc"

    config = build_contract_config(fields, args.template, output_name)

    # Step 4: 保存config（可选）
    if args.save_config:
        config_path = os.path.join(tempfile.gettempdir(), 'contract_config.json')
        with open(config_path, 'w', encoding='utf-8') as f:
            json.dump(config, f, ensure_ascii=False, indent=2)
        print(f"\n[*] 配置已保存: {config_path}")

    # Step 5: 只显示结果（dry-run模式）
    if args.dry_run:
        print("\n[*] 干运行模式，不填入合同")
        print("\n将要填入的字段:")
        for k, v in config['fields'].items():
            print(f"  {k} -> {v}")
        return

    # Step 6: 填入合同
    print("\n" + "=" * 60)
    print("  填入合同")
    print("=" * 60)
    fill_contract(config)


if __name__ == '__main__':
    main()
