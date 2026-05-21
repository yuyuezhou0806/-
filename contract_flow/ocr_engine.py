import re
import tempfile
import os
from paddleocr import PaddleOCR

_ocr = None


def get_ocr():
    global _ocr
    if _ocr is None:
        _ocr = PaddleOCR(
            use_angle_cls=True,
            lang="ch",
            show_log=False,
            use_gpu=False
        )
    return _ocr


def extract_field(text_lines, keywords):
    """从文本行中提取关键词后面的内容，过滤无效值"""
    INVALID_VALUES = {"委托单位", "（委托单位）", "（检测机构）", "乙方", "检测机构"}
    for line in text_lines:
        for kw in keywords:
            if kw in line:
                # 尝试匹配 "关键词: 内容" 或 "关键词：内容"
                pattern = rf"{re.escape(kw)}[：:]\s*(.+)"
                m = re.search(pattern, line)
                if m:
                    val = m.group(1).strip()
                    if val and val not in INVALID_VALUES and len(val) > 1:
                        return val
                # 如果没有冒号，尝试取关键词后面的内容
                idx = line.find(kw)
                if idx >= 0:
                    after = line[idx + len(kw):].strip()
                    # 清理前导符号
                    after = re.sub(r"^[：:\s()（）]+", "", after)
                    if after and after not in INVALID_VALUES and len(after) > 1:
                        return after
    return ""


def parse_contract_ocr(image_bytes):
    """识别合同图片，提取关键字段"""
    ocr = get_ocr()

    with tempfile.NamedTemporaryFile(suffix=".jpg", delete=False) as tmp:
        tmp.write(image_bytes)
        tmp_path = tmp.name

    try:
        result = ocr.ocr(tmp_path, cls=True)
    finally:
        os.unlink(tmp_path)

    if not result or not result[0]:
        return {"project_name": "", "client_unit": "", "payer_unit": ""}

    text_lines = []
    for line in result[0]:
        if line:
            text = line[1][0]
            text_lines.append(text)

    return _extract_from_lines(text_lines)


def parse_pdf_file(file_bytes):
    """解析 PDF 文件，提取文本内容"""
    try:
        import fitz
    except ImportError:
        raise ImportError("解析 PDF 需要安装 PyMuPDF，请运行: pip install PyMuPDF")

    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
        tmp.write(file_bytes)
        tmp_path = tmp.name

    text_lines = []
    try:
        doc = fitz.open(tmp_path)
        for page in doc:
            text = page.get_text()
            for line in text.splitlines():
                line = line.strip()
                if line:
                    text_lines.append(line)
        doc.close()
    finally:
        os.unlink(tmp_path)

    return _extract_from_lines(text_lines)


def parse_word_file(file_bytes):
    """解析 Word 文件，提取文本内容"""
    try:
        import docx
    except ImportError:
        raise ImportError("解析 Word 需要安装 python-docx，请运行: pip install python-docx")

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp.write(file_bytes)
        tmp_path = tmp.name

    text_lines = []
    try:
        document = docx.Document(tmp_path)
        for para in document.paragraphs:
            line = para.text.strip()
            if line:
                text_lines.append(line)
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        line = para.text.strip()
                        if line:
                            text_lines.append(line)
    finally:
        os.unlink(tmp_path)

    return _extract_from_lines(text_lines)


def parse_doc_file(file_bytes):
    """解析 .doc 文件（需要系统安装 Microsoft Word 或 pywin32）"""
    try:
        import win32com.client
    except ImportError:
        raise ImportError("解析 .doc 需要安装 pywin32，请运行: pip install pywin32")

    with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as tmp:
        tmp.write(file_bytes)
        tmp_path = tmp.name

    word = None
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        word.DisplayAlerts = False
        doc = word.Documents.Open(tmp_path)
        text = doc.Content.Text
        doc.Close()
        text_lines = [line.strip() for line in text.splitlines() if line.strip()]
        return _extract_from_lines(text_lines)
    except Exception as e:
        raise ValueError(f"Word 解析失败: {str(e)}")
    finally:
        if word:
            try:
                word.Quit()
            except Exception:
                pass
        if os.path.exists(tmp_path):
            os.unlink(tmp_path)


def _clean_value(val):
    """清理提取到的值，去掉无效后缀"""
    if not val:
        return ""
    val = val.strip()
    # 纯无效值
    if val in {"（盖章）", "(盖章)", "盖章", "无", "有", "", "：", ":"}:
        return ""
    # 去掉末尾的（盖章）
    val = re.sub(r"[（(]盖章[）)]\s*$", "", val).strip()
    # 去掉末尾冒号
    val = re.sub(r"[：:]\s*$", "", val).strip()
    return val


def _is_valid_unit(val):
    """判断值是否是有效的单位/公司名称，过滤掉条款描述等无效内容"""
    if not val or len(val) < 4:
        return False
    # 必须包含组织名称特征词
    org_keywords = ["公司", "集团", "中心", "研究院", "研究所", "大学", "学院"]
    if not any(kw in val for kw in org_keywords):
        return False
    # 过滤明显是描述性文本的噪声词
    noise = ["委托的检测", "能力等级", "根据", "规定", "签订", "双方", "遵循", "甲乙", "之权利", "之义务"]
    for n in noise:
        if n in val:
            return False
    return True


def _extract_from_lines(text_lines):
    """从文本行中提取合同关键字段"""
    full_text = "\n".join(text_lines)
    continuous_text = re.sub(r'\s+', ' ', full_text)

    # ===== 提取工程名称 =====
    project_name = ""
    for i, line in enumerate(text_lines):
        if "工程名称" in line:
            # 尝试从当前行提取（支持冒号后内容）
            m = re.search(r"工程名称\s*[：:]\s*(.+)", line)
            if m:
                val = _clean_value(m.group(1))
                if val:
                    project_name = val
                    break
            # 当前行只有关键词，值在下一行
            if i + 1 < len(text_lines):
                next_val = _clean_value(text_lines[i + 1])
                if next_val and len(next_val) > 3:
                    project_name = next_val
                    break
    # 连续文本兜底
    if not project_name:
        m = re.search(r"工程名称\s*[：:]\s*([^：:\n]+?)(?=\s*(工程地址|甲方|乙方|$|\n))", continuous_text)
        if m:
            project_name = _clean_value(m.group(1))

    # ===== 提取委托单位（甲方） =====
    client_unit = ""

    # 方法1：逐行精确匹配，优先匹配"甲方"
    keywords_priority = ["甲方（委托单位）", "甲方(委托单位)", "甲方", "委托单位", "建设单位"]
    for i, line in enumerate(text_lines):
        matched = False
        for kw in keywords_priority:
            if kw in line:
                matched = True
                # 尝试从当前行提取值（支持中文冒号、英文冒号、空格）
                m = re.search(rf"{re.escape(kw)}\s*[：:]\s*(.+)", line)
                if m:
                    val = _clean_value(m.group(1))
                    if val and _is_valid_unit(val):
                        client_unit = val
                        break
                # 没有冒号的情况，关键词后直接跟值
                idx = line.find(kw)
                after = line[idx + len(kw):].strip()
                after = re.sub(r"^[：:\s()（）]+", "", after)
                val = _clean_value(after)
                if val and _is_valid_unit(val):
                    client_unit = val
                    break
        if client_unit:
            break

        # 当前行只有关键词没有值，看下一行
        if matched and i + 1 < len(text_lines):
            next_val = _clean_value(text_lines[i + 1])
            if next_val and _is_valid_unit(next_val) and next_val != project_name:
                client_unit = next_val
                break

    # 方法2：连续文本正则兜底（同样要过有效性校验）
    if not client_unit:
        patterns = [
            r"甲方[（(]委托单位[）)]\s*[：:]\s*([^：:\n]+?)(?=\s*(乙方|工程名称|工程地址|$|\n))",
            r"甲方\s*[：:]\s*([^：:\n]+?)(?=\s*(乙方|工程名称|工程地址|$|\n))",
            r"委托单位\s*[：:]\s*([^：:\n]+?)(?=\s*(乙方|工程名称|工程地址|$|\n))",
        ]
        for pattern in patterns:
            m = re.search(pattern, continuous_text)
            if m:
                val = _clean_value(m.group(1))
                if val and _is_valid_unit(val):
                    client_unit = val
                    break

    # 方法3：在"甲方"附近搜索公司名称（兜底）
    if not client_unit:
        jia_idx = continuous_text.find("甲方")
        if jia_idx >= 0:
            nearby = continuous_text[jia_idx:jia_idx + 300]
            # 优先匹配"甲方"后面直到换行或特定分隔符的内容
            m = re.search(r"甲方\s*[：:]\s*([^\n\r]+?)(?:\n|\r|乙方|工程名称|$)", nearby)
            if m:
                val = _clean_value(m.group(1))
                if val and _is_valid_unit(val):
                    client_unit = val
            if not client_unit:
                m = re.search(r"([^：:\s]{2,}?(?:有限公司|有限责任公司|股份公司|集团|中心|研究院|研究所))", nearby)
                if m:
                    val = _clean_value(m.group(1))
                    if val and _is_valid_unit(val):
                        client_unit = val

    # ===== 付款单位 = 委托单位 =====
    payer_unit = client_unit

    return {
        "project_name": project_name,
        "client_unit": client_unit,
        "payer_unit": payer_unit
    }


def parse_contract_file(file_bytes, filename):
    """根据文件类型解析合同文件，提取关键字段"""
    lower_name = filename.lower()

    if lower_name.endswith(".pdf"):
        return parse_pdf_file(file_bytes)
    elif lower_name.endswith(".docx"):
        return parse_word_file(file_bytes)
    elif lower_name.endswith(".doc"):
        return parse_doc_file(file_bytes)
    elif lower_name.endswith((".jpg", ".jpeg", ".png", ".bmp", ".tiff", ".tif", ".webp")):
        return parse_contract_ocr(file_bytes)
    else:
        raise ValueError(f"不支持的文件格式: {filename}")
