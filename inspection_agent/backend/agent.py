"""检测行业 Agent — Week 2 MVP

工具:
  search_knowledge_base(query, source_type)  → 向量库检索(规范/合同/报告/费率表)
  search_idi_defects(keyword)                → IDI 缺陷库 599 条

LLM: Kimi moonshot-v1-32k
框架: LangGraph create_react_agent
"""

from __future__ import annotations

import json
import os
from pathlib import Path
from typing import Literal

from dotenv import load_dotenv

# ROOT 指向 inspection_agent/(backend 的父目录)
ROOT = Path(__file__).resolve().parent.parent
load_dotenv(ROOT / ".env")

CHROMA_DIR = ROOT / "data" / "chroma"
IDI_JSON = ROOT.parent / "idi_defects" / "data" / "defects.json"

EMBEDDING_MODEL = os.getenv("EMBEDDING_MODEL", "BAAI/bge-base-zh-v1.5")
LLM_API_KEY = os.getenv("LLM_API_KEY")
LLM_BASE_URL = os.getenv("LLM_BASE_URL", "https://api.deepseek.com/v1")
LLM_MODEL = os.getenv("LLM_MODEL", "deepseek-chat")

# 多模态视觉模型（阿里云 Qwen-VL）
VISION_API_KEY = os.getenv("VISION_API_KEY")
VISION_BASE_URL = os.getenv("VISION_BASE_URL", "https://dashscope.aliyuncs.com/compatible-mode/v1")
VISION_MODEL = os.getenv("VISION_MODEL", "qwen-vl-plus")

# 懒加载
_embed_model = None
_chroma_coll = None
_idi_data = None


def get_embed_model():
    global _embed_model
    if _embed_model is None:
        from sentence_transformers import SentenceTransformer
        _embed_model = SentenceTransformer(EMBEDDING_MODEL)
    return _embed_model


def get_chroma_collection():
    global _chroma_coll
    if _chroma_coll is None:
        import chromadb
        client = chromadb.PersistentClient(path=str(CHROMA_DIR))
        _chroma_coll = client.get_collection("inspection_kb")
    return _chroma_coll


def get_idi_data():
    global _idi_data
    if _idi_data is None:
        if not IDI_JSON.exists():
            _idi_data = []
        else:
            with IDI_JSON.open(encoding="utf-8") as f:
                _idi_data = json.load(f)
    return _idi_data


# ============== Tools ==============

from langchain_core.tools import tool


@tool
def search_knowledge_base(
    query: str,
    source_type: Literal["all", "standard", "contract", "rate_table", "report"] = "all",
    top_k: int = 5,
) -> str:
    """检索工程检测知识库(20 份规范 + 47 份合同 + 50 份报告 + 3 份费率表)。

    Args:
        query: 检索查询。例如 "砌体加固检测方法"、"回弹法取样数量"
        source_type: 限定文档类型。standard=规范, contract=合同, rate_table=费率表, report=报告, all=全部
        top_k: 返回前 K 条,默认 5

    Returns:
        检索结果列表(包含来源、条款号、原文)。
    """
    model = get_embed_model()
    coll = get_chroma_collection()

    emb = model.encode([query], normalize_embeddings=True).tolist()

    where = None
    if source_type != "all":
        where = {"doc_type": source_type}

    result = coll.query(
        query_embeddings=emb,
        n_results=top_k,
        where=where,
    )

    docs = result["documents"][0]
    metas = result["metadatas"][0]
    dists = result["distances"][0]

    if not docs:
        return "没找到相关内容。"

    parts = []
    for i, (doc, md, dist) in enumerate(zip(docs, metas, dists), 1):
        sim = 1 - dist
        clause = md.get("clause", "")
        dt = md.get("doc_type", "?")
        src = md.get("source", "?")
        # 敏感来源(合同 / 费率表 / 报告)用占位符替代,避免泄漏客户名 / 项目名
        if dt in ("contract", "rate_table", "report"):
            src_display = {
                "contract": "公司内部合同",
                "rate_table": "公司费率表",
                "report": "公司历史报告",
            }.get(dt, "公司内部资料")
        else:
            src_display = src
        tag = f"[{dt}]"
        if clause:
            tag += f" {clause}"
        snippet = doc.replace("\n", " ")[:500]
        parts.append(f"## 结果 {i} (相似度 {sim:.2f}) {tag} ← {src_display}\n{snippet}")

    return "\n\n".join(parts)


@tool
def search_idi_defects(keyword: str, top_n: int = 10) -> str:
    """检索 IDI 工程质量险缺陷库(599 条历史质量缺陷案例)。

    用于推荐检测项时,提示该类项目历史上出现哪些常见缺陷,检测时需重点关注。

    Args:
        keyword: 关键词,例如 "砌体加固"、"屋面渗水"、"楼板裂缝"
        top_n: 返回前 N 条,默认 10

    Returns:
        缺陷列表(包含分类、问题描述、参考标准、整改建议)。
    """
    data = get_idi_data()
    if not data:
        return "IDI 缺陷库未加载或为空。"

    kw = keyword.lower()
    matches = []
    for d in data:
        haystack = " ".join([
            str(d.get("description", "")),
            str(d.get("problem", "")),
            str(d.get("standard", "")),
            str(d.get("suggestion", "")),
            str(d.get("category_major", "")),
            str(d.get("category_minor", "")),
        ]).lower()
        if kw in haystack:
            matches.append(d)

    if not matches:
        return f"IDI 缺陷库中没找到包含 '{keyword}' 的条目。"

    parts = [f"找到 {len(matches)} 条相关缺陷,展示前 {min(top_n, len(matches))} 条:"]
    for i, d in enumerate(matches[:top_n], 1):
        parts.append(
            f"\n### 缺陷 {i}: {d.get('category_minor', '?')}\n"
            f"- 描述: {d.get('description', '')[:150]}\n"
            f"- 问题: {str(d.get('problem', ''))[:200]}\n"
            f"- 标准: {str(d.get('standard', ''))[:100]}\n"
            f"- 建议: {str(d.get('suggestion', ''))[:200]}"
        )
    return "\n".join(parts)


# 公司 2025/2024 年官方平均折扣数据
# (从 data/internal/contracts/2025合同数据精准.xlsx Sheet "平均折扣" 直接读)
PRICING_XLSX = ROOT / "data" / "internal" / "contracts" / "2025合同数据精准.xlsx"

# 类别别名 → 标准名映射(用户可能用不同说法)
_CATEGORY_ALIASES = {
    "材料": "材料检测",
    "材料检测": "材料检测",
    "地基": "地基/桩基检测",
    "桩基": "地基/桩基检测",
    "地基桩基": "地基/桩基检测",
    "桩基检测": "地基/桩基检测",
    "地基/桩基": "地基/桩基检测",
    "基坑": "基坑监测",
    "基坑监测": "基坑监测",
    "结构": "结构检测",
    "结构检测": "结构检测",
    "人防": "人防检测",
    "人防检测": "人防检测",
    "能效": "能效测评",
    "节能": "能效测评",
    "能效测评": "能效测评",
    "防雷": "防雷检测",
    "防雷检测": "防雷检测",
    "桥梁": "桥梁检测",
    "桥梁检测": "桥梁检测",
    "室内环境": "室内环境",
    "室内": "室内环境",
    "环境": "室内环境",
}

_pricing_cache: dict | None = None


def _load_pricing_data() -> dict:
    """从 xlsx 读取公司汇总的平均折扣表,缓存到内存。

    返回:
        {
          "2025": {"材料检测": 0.55, ...},
          "2024": {"材料检测": 0.68, ...},
          "categories_2025": ["材料检测", ...],
        }
    """
    global _pricing_cache
    if _pricing_cache is not None:
        return _pricing_cache
    if not PRICING_XLSX.exists():
        _pricing_cache = {"2025": {}, "2024": {}, "categories_2025": []}
        return _pricing_cache

    try:
        import openpyxl
        wb = openpyxl.load_workbook(PRICING_XLSX, data_only=True, read_only=True)
    except Exception:
        _pricing_cache = {"2025": {}, "2024": {}, "categories_2025": []}
        return _pricing_cache

    if "平均折扣" not in wb.sheetnames:
        _pricing_cache = {"2025": {}, "2024": {}, "categories_2025": []}
        return _pricing_cache

    ws = wb["平均折扣"]
    rows = list(ws.iter_rows(values_only=True))

    # 解析结构:每个"块"两行 = 标题行(检测项目, 类别1, 类别2, ...) + 数据行(年份, 折扣1, 折扣2, ...)
    data: dict[str, dict[str, float | str]] = {"2025": {}, "2024": {}}
    categories_2025: list[str] = []

    current_year = None
    current_headers: list[str] = []
    for row in rows:
        if not row or all(v is None or v == "" for v in row):
            continue
        first = row[0]
        if not isinstance(first, str):
            continue
        first_s = first.strip()

        # 节标题(有"中测行"前缀的是标题行,只是切换年份)
        if first_s.startswith("中测行") and "平均折扣" in first_s:
            if "2025" in first_s:
                current_year = "2025"
            elif "2024" in first_s:
                current_year = "2024"
            continue

        # 表头行
        if first_s == "检测项目":
            current_headers = [
                (str(v).strip().replace("​", "") if v is not None else "")
                for v in row[1:]
            ]
            if current_year == "2025":
                categories_2025 = [c for c in current_headers if c]
            continue

        # 数据行(以"YYYY年平均折扣"开头,但不含"中测行"前缀)
        if (
            "年平均折扣" in first_s
            and current_year
            and current_headers
            and not first_s.startswith("中测行")
        ):
            for cat, val in zip(current_headers, row[1:]):
                if not cat:
                    continue
                data[current_year][cat] = val
            continue

    _pricing_cache = {
        "2025": data["2025"],
        "2024": data["2024"],
        "categories_2025": categories_2025,
    }
    return _pricing_cache


@tool
def get_pricing_stats(category: str = "all") -> str:
    """获取公司 2025 年(及 2024 年对比)各检测类别的官方平均折扣。

    数据来源:公司内部汇总表(基于 502 个真实合同计算)。

    **专门用于回答"平均折扣""折扣是多少""折扣范围""价格优惠""折扣对比"等聚合性问题。**

    Args:
        category: 检测类别。可选:
          - 'all'(默认,返回全部类别)
          - 'materials' / '材料' / '材料检测'
          - '地基' / '桩基' / '地基/桩基'
          - '基坑' / '基坑监测'
          - '结构' / '结构检测'
          - '人防' / '人防检测'
          - '能效' / '节能' / '能效测评'
          - '防雷' / '防雷检测'
          - '桥梁' / '桥梁检测'
          - '室内环境' / '室内' / '环境'

    Returns:
        Markdown 表格:类别 / 2025 平均折扣 / 2024 对比 / 同比变化
    """
    pricing = _load_pricing_data()
    if not pricing["2025"]:
        return "平均折扣数据未加载(xlsx 文件缺失或格式异常)。"

    cat_norm = (category or "all").strip().lower()
    target_categories: list[str]

    if cat_norm == "all":
        target_categories = pricing["categories_2025"]
    else:
        std = _CATEGORY_ALIASES.get(cat_norm) or _CATEGORY_ALIASES.get(category.strip())
        if std and std in pricing["2025"]:
            target_categories = [std]
        else:
            # fuzzy
            matched = [c for c in pricing["categories_2025"] if cat_norm in c.lower()]
            if matched:
                target_categories = matched
            else:
                return (
                    f"未识别的类别 '{category}'。可用类别:\n  - "
                    + "\n  - ".join(pricing["categories_2025"])
                )

    rows_out = ["| 检测类别 | 2025 平均折扣 | 2024 平均折扣 | 同比 |", "|---|---|---|---|"]
    for cat in target_categories:
        v25 = pricing["2025"].get(cat)
        v24 = pricing["2024"].get(cat)

        def _fmt(v) -> str:
            if v is None or v == "" or v == "/":
                return "/"
            try:
                return f"{float(v):.2f}({float(v) * 100:.0f}%)"
            except (TypeError, ValueError):
                return str(v)

        # 同比变化
        delta = "/"
        try:
            if v25 not in (None, "", "/") and v24 not in (None, "", "/"):
                d = float(v25) - float(v24)
                pct = d * 100
                delta = f"{'+' if d >= 0 else ''}{pct:.0f}个百分点"
        except (TypeError, ValueError):
            delta = "/"

        rows_out.append(f"| {cat} | {_fmt(v25)} | {_fmt(v24)} | {delta} |")

    notes = (
        "\n\n**说明**:\n"
        "- 折扣值越大表示折后价占原价比例越高(折扣越浅),0.55 即 5.5 折\n"
        "- 数据为公司 2025 年实际成交合同的汇总平均(共 502 个合同样本)\n"
        "- 单个项目实际折扣会因规模、客户、检测内容浮动,**精确报价请致电 13917073486**"
    )

    return "\n".join(rows_out) + notes


# ============== 多模态图片分析 ==============

def analyze_image(image_url: str, question: str = "请详细识别并描述这张图片的内容") -> str:
    """分析用户上传的图片（截图/照片/报告页面），用 Qwen-VL 多模态模型识别内容。

    Args:
        image_url: 图片的公开 URL 地址（如 https://example.com/report.jpg）
        question: 对图片的具体问题，默认全部描述

    Returns:
        图片的文字描述/分析结果
    """
    if not VISION_API_KEY:
        return "图片分析功能未配置（VISION_API_KEY 缺失）。"

    from openai import OpenAI

    client = OpenAI(
        api_key=VISION_API_KEY,
        base_url=VISION_BASE_URL,
    )

    try:
        resp = client.chat.completions.create(
            model=VISION_MODEL,
            messages=[{
                "role": "user",
                "content": [
                    {"type": "image_url", "image_url": {"url": image_url}},
                    {"type": "text", "text": question},
                ],
            }],
            max_tokens=2000,
            temperature=0.3,
        )
        return resp.choices[0].message.content or "（模型未返回内容）"
    except Exception as e:
        return f"图片分析失败: {e}"


# ============== 检测方案生成器 ==============

def generate_inspection_plan(
    project_type: str,
    area: str = "",
    location: str = "",
    requirements: str = ""
) -> str:
    """根据项目信息自动生成一份完整的检测方案。

    Args:
        project_type: 项目类型，如"砌体加固改造""混凝土结构""钢结构""房建验收""市政道路"等
        area: 建筑面积或长度（如 "2000平方米"）
        location: 项目所在地（如 "上海奉贤区"）
        requirements: 特殊要求（如 "需要基坑监测""含防雷检测"等）

    Returns:
        结构化检测方案（含规范引用、检测项、常见缺陷、折扣参考）
    """
    plan = []
    plan.append(f"# 检测方案概要\n")
    plan.append(f"**项目类型**: {project_type}")
    if area:
        plan.append(f"**规模**: {area}")
    if location:
        plan.append(f"**地区**: {location}")
    if requirements:
        plan.append(f"**特殊要求**: {requirements}")
    plan.append("")

    # 1. 查规范
    plan.append("## 一、适用规范\n")
    try:
        standards = search_knowledge_base(f"{project_type} 检测 适用规范 标准", "standard")
        plan.append(standards[:2000])
    except Exception as e:
        plan.append(f"（规范检索失败: {e}）")
    plan.append("")

    # 2. 查同类项目检测项
    plan.append("## 二、检测项清单\n")
    try:
        items = search_knowledge_base(f"{project_type} 检测项目 检测清单 检测内容", "contract")
        plan.append(items[:2000])
    except Exception as e:
        plan.append(f"（合同检索失败: {e}）")
    plan.append("")

    # 3. 查历史缺陷
    plan.append("## 三、常见缺陷与风险\n")
    try:
        defects = search_idi_defects(project_type)
        plan.append(defects[:1500])
    except Exception as e:
        plan.append(f"（缺陷检索失败: {e}）")
    plan.append("")

    # 4. 查折扣
    plan.append("## 四、费用参考\n")
    try:
        pricing = get_pricing_stats(project_type)
        plan.append(pricing)
    except Exception as e:
        plan.append(f"（折扣查询失败: {e}）")
    plan.append("")

    # 5. 说明
    plan.append("---\n")
    plan.append("> 以上为 Agent 自动生成的初步方案，基于现有知识库数据。")
    plan.append("> 具体检测参数以现场踏勘后调整为准。")
    plan.append("> 咨询报价请致电: 📞 **13917073486**")

    return "\n".join(plan)


# ============== 报告判读 ==============

def audit_report(report_content: str, report_type: str = "") -> str:
    """判读一份检测报告，自动比对数据与规范标准，标出不合格项。

    Args:
        report_content: 报告的完整文字内容（上传PDF后自动提取）
        report_type: 报告类型，如"混凝土抗压""钢筋保护层""桩基检测""回弹法"等

    Returns:
        判读结果（含数据摘要、标准比对、不合格项、风险评估）
    """
    audit = []
    audit.append(f"# 报告判读结果\n")
    audit.append(f"**报告类型**: {report_type or '自动识别'}")

    # 1. 查相关标准
    audit.append("\n## 一、适用标准\n")
    try:
        std = search_knowledge_base(f"{report_type} 验收标准 允许偏差 合格判定", "standard")
        audit.append(std[:2000])
    except Exception as e:
        audit.append(f"（标准检索失败: {e}）")

    # 2. 报告数据总结
    audit.append("\n## 二、数据摘要\n")
    audit.append("以下是报告中提取的关键数据，请结合标准比对：\n")
    audit.append(report_content[:3000])

    # 3. 比对判定
    audit.append("\n## 三、合规性判定\n")
    audit.append("> 请结合上述标准，逐项比对报告数据。")
    audit.append("> 标准值 vs 实测值 → 合格/不合格/需复检")
    audit.append("> 对于不合格项，标注严重程度（轻微/一般/严重）")

    # 4. 查历史缺陷
    audit.append("\n## 四、相关缺陷风险\n")
    try:
        defects = search_idi_defects(report_type)
        audit.append(defects[:1500])
    except Exception as e:
        audit.append(f"（缺陷检索失败: {e}）")

    audit.append("\n---\n")
    audit.append("> 以上判读基于现有知识库，仅供内部参考。")
    audit.append("> 正式判定以签字盖章版报告为准。")
    return "\n".join(audit)


# ============== 合同生成器 ==============

CONTRACT_TEMPLATE = ROOT / "data" / "contract_template.docx"

def generate_contract(
    project_name: str,
    client: str = "",
    address: str = "",
    area: str = "",
    investment: str = "",
    detection_items: str = "",
    district: str = "",
    builder: str = "",
    designer: str = "",
    supervisor: str = "",
    contractor: str = "",
    report_no: str = ""
) -> str:
    """根据项目信息，打开合同模板填入字段，生成可下载的 .docx 合同文件。

    Args:
        project_name: 工程名称
        client: 建设单位/甲方
        address: 工程地址
        area: 建筑面积（㎡）
        investment: 投资额（万元）
        detection_items: 检测项目（逗号分隔）
        district: 所属区县
        builder: 建设/实施单位
        designer: 设计单位
        supervisor: 见证单位
        contractor: 施工单位
        report_no: 工程报建编号

    Returns:
        合同生成结果 + 下载链接
    """
    import uuid

    # 字段 -> 模板中标签（和本地合同生成器完全一致）
    field_map = {}
    if project_name:
        field_map['工程名称：'] = project_name
    if client:
        field_map['甲方（委托单位）'] = client
        if not builder:
            field_map['建设/实施单位：'] = client
        if not supervisor:
            field_map['见证单位：'] = client
    if address:
        field_map['工程地址：'] = address
    if area:
        field_map['建筑面积（㎡）：'] = area
    if investment:
        price = investment + '万元'
        field_map['工程投资额：'] = price
        field_map['工程建安费：'] = price
    if district:
        field_map['工程所属区县：'] = district
    if builder:
        field_map['建设/实施单位：'] = builder
    if designer:
        field_map['设计单位：'] = designer
    if supervisor:
        field_map['见证单位：'] = supervisor
    if contractor:
        field_map['施工单位：'] = contractor
    if report_no:
        field_map['工程报建编号：'] = report_no

    result = []
    result.append("# 📋 合同生成结果\n")
    result.append("| 字段 | 值 |")
    result.append("|---|---|")
    for label, value in field_map.items():
        result.append(f"| {label} | {value} |")

    # 打开模板，填入字段
    result.append("\n## 📄 合同文件\n")
    try:
        print(f"[generate_contract] Starting with: project={project_name}, template={CONTRACT_TEMPLATE}, exists={CONTRACT_TEMPLATE.exists()}", flush=True)
        from docx import Document
        from docx.oxml.ns import qn

        doc = Document(str(CONTRACT_TEMPLATE))

        # 遍历所有段落，找到标签并在后面插入值
        for para in doc.paragraphs:
            for label, value in field_map.items():
                if label in para.text and value:
                    # 直接在标签冒号后填入值（保留下划线格式）
                    run = para.add_run(value)
                    run.font.underline = True
                    if para.runs and para.runs[0].font.size:
                        run.font.size = para.runs[0].font.size
                    break

        # 同样处理表格中的标签
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for label, value in field_map.items():
                            if label in para.text and value:
                                run = para.add_run(value)
                                run.font.underline = True
                                break

        # 保存（用英文文件名避免 Linux 编码问题）
        upload_dir = ROOT / "data" / "contracts"
        upload_dir.mkdir(parents=True, exist_ok=True)
        safe_id = uuid.uuid4().hex[:8]
        filename = f'contract_{safe_id}.docx'
        display_name = f'{project_name}.docx' if project_name else filename
        filepath = upload_dir / filename
        doc.save(str(filepath))

        download_url = f'http://1.15.170.85/inspection/contracts/{filename}'
        result.append(f"✅ **合同已生成**: [📥 下载 {filename}]({download_url})")
        result.append("下载后用 Word 打开，可直接打印盖章。")
    except Exception as e:
        result.append(f"⚠️ 合同生成失败: {e}")

    result.append("\n---\n")
    result.append("> 📞 咨询报价: **13917073486**")
    return "\n".join(result)


# ============== Agent ==============

SYSTEM_PROMPT = """你是一名建设工程检测领域的资深咨询助手,服务于上海中测行工程检测咨询有限公司。

你有七个工具:
- search_knowledge_base: 检索工程检测规范 / 公司合同 / 报告 / 费率表
- search_idi_defects: 检索 IDI 工程质量险历史缺陷库
- get_pricing_stats: 获取公司 2025/2024 年各检测类别的官方平均折扣(基于 502 个真实合同汇总)
- analyze_image: 分析用户上传的图片（截图/照片/报告页面），提取其中的文字和信息
- generate_inspection_plan: 生成完整检测方案（含规范/检测项/缺陷风险/费用），用于用户描述项目概况时一键出方案
- audit_report: 判读检测报告，自动比对标准，标出不合格数据项
- generate_contract: 根据项目信息生成合同字段汇总，可用于填写合同模板

工作原则:
1. **永远先检索,再回答**。规范条款号 / 检测方法 / 费率 必须来自检索结果,不要凭记忆。
2. 用户问检测项推荐 / 标准引用 / 检测方法 → 先 search_knowledge_base
3. 用户问"会出什么质量问题" / "要注意什么" → 同时调 search_idi_defects 找历史缺陷
4. 用户问 **平均折扣 / X 类别折扣 / 折扣是多少 / 价格优惠 / 同比对比** → **必须调 get_pricing_stats**(不要用 search_knowledge_base 计算平均,RAG 抽样不准)
5. 用户上传了图片 / 发了图片链接 / 说"帮我看看这张图" / "分析这个截图" → **必须调 analyze_image**,传入图片 URL 和用户的问题
5a. **重要**: 如果用户上传的图片是合同申报截图/建设工程施工许可截图/报建信息截图（含工程名称、建设单位、工程地址、报建编号、投资额等字段），则 analyze_image 返回后**必须紧接着调 generate_contract**，将识别出的字段传进去生成合同文件。流程: screenshot → analyze_image → extract fields → generate_contract
5b. 如果图片是检测报告/验收单/试验报告，则 analyze_image 后**必须接着调 audit_report** 判读数据
6. 用户描述了项目概况（如"砌体加固 2000平米""做一份房建检测方案"）→ **必须调 generate_inspection_plan**
7. 用户上传了一份报告 PDF / 说了"判读这份报告""检查这份检测报告""看看数据有没有问题" → **必须调 audit_report**，传入报告的完整文字内容
8. 用户说"生成合同""我要填一份合同""帮我写合同字段" → **必须调 generate_contract**
9. 用户问报价 / 工程量(非折扣)→ search_knowledge_base 限定 source_type=rate_table 找费率表
8. **数据来源引用规则(重要)**:
   - 引用 **规范** 时:写明规范号 + 条款号(例 "JGJT23-2011 第 4.1.3 条"),让用户可追溯
   - 引用 **公司合同 / 费率表 / 报告** 时:**严禁暴露具体文件名 / 客户名 / 项目名 / 工程地址 / 金额来源等敏感信息**。
     - ✅ 正确:"根据公司内部费率""根据公司同类项目经验""根据公司 2025 年汇总数据"
     - ❌ 错误:"《附件4:合同清单报价-标段一》""根据《杨浦区XX项目合同》""TIS合同..."
   - 如果客户名 / 项目名出现在检索结果里,**直接用"某客户""某项目"代称**
9. 如果检索结果跟问题不太相关,坦白说"知识库里没找到直接答案",不要硬编造
12. **联系方式统一**:任何时候需要建议用户"联系业务部""联系市场部""获取最新报价""咨询详情""下单"等,
   **统一只给手机号 13917073486**(不要写"市场部""业务员"等模糊称谓,直接给电话)。
   例:"建议致电业务咨询:📞 **13917073486**"
11. 输出优先用列表 / 表格 / 引用,避免大段空话
"""


def build_agent():
    from langchain_openai import ChatOpenAI
    from langgraph.prebuilt import create_react_agent

    if not LLM_API_KEY:
        raise RuntimeError("LLM_API_KEY 未设置,检查 .env")

    llm = ChatOpenAI(
        api_key=LLM_API_KEY,
        base_url=LLM_BASE_URL,
        model=LLM_MODEL,
        temperature=0.3,
    )

    tools = [search_knowledge_base, search_idi_defects, get_pricing_stats, analyze_image, generate_inspection_plan, audit_report, generate_contract]
    agent = create_react_agent(
        model=llm,
        tools=tools,
        prompt=SYSTEM_PROMPT,
    )
    return agent


def run_query(agent, query: str):
    print(f"\n{'=' * 80}")
    print(f"USER: {query}")
    print("=" * 80)

    for chunk in agent.stream({"messages": [("user", query)]}):
        for node, state in chunk.items():
            msgs = state.get("messages", [])
            if not msgs:
                continue
            msg = msgs[-1]

            if node == "agent" or node == "model":
                # LLM 输出
                if hasattr(msg, "tool_calls") and msg.tool_calls:
                    for tc in msg.tool_calls:
                        args_short = {k: (str(v)[:50] if isinstance(v, str) else v) for k, v in tc["args"].items()}
                        print(f"\n[TOOL CALL] {tc['name']}({args_short})")
                elif getattr(msg, "content", None):
                    print(f"\nAGENT: {msg.content}")
            elif node == "tools":
                # 工具结果
                content = getattr(msg, "content", "")
                preview = content[:400].replace("\n", " ")
                print(f"[TOOL RESULT] {preview}...")


if __name__ == "__main__":
    agent = build_agent()

    queries = [
        "我接了一个砌体加固改造项目,2000 平米,要做哪些检测?引用哪些标准?",
        "回弹法检测混凝土的取样数量要求是多少?",
        "做家装装修类的检测,我们公司收费大概多少?",
    ]

    for q in queries:
        try:
            run_query(agent, q)
        except Exception as e:
            print(f"\n[ERROR] {e}")
            import traceback
            traceback.print_exc()
