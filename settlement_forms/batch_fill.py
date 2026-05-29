#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""批量填充项目结算报审单"""

import os
import shutil
from datetime import datetime
from docx import Document
import pandas as pd


def format_amount(value):
    """Excel 中的数值是万元，转成元。保留两位小数（如有），去掉末尾的 .00。"""
    if pd.isna(value):
        return ""
    v = float(value) * 10000
    if v == int(v):
        return f"{int(v)}"
    return f"{v:.2f}"


def fill_paragraph(paragraph, new_text):
    """清空段落所有 runs，把新文本写入第一个 run，保留格式属性。"""
    first_run = None
    for run in paragraph.runs:
        if first_run is None:
            first_run = run
        run.text = ""
    if first_run:
        first_run.text = new_text
    else:
        paragraph.add_run(new_text)


def fill_settlement_form(template_path, output_path, data):
    """根据单条数据填充模板并保存。"""
    shutil.copy(template_path, output_path)
    doc = Document(output_path)

    # 金额字段：万元 → 元
    total_workload = format_amount(data.get("实际发生总工作量(含税)", ""))
    test_fee = format_amount(data.get("实际应支付检测费（含税）", ""))
    confirmed_income = format_amount(data.get("已确认收入（含税）", ""))
    invoice_amount = format_amount(data.get("开票金额(含税)", ""))
    received_amount = format_amount(data.get("已收款金额（含税）", ""))

    def _get_str(key):
        val = data.get(key, "")
        if pd.isna(val):
            return ""
        return str(val).strip()

    # 其他字段
    payer = _get_str("付款单位")
    contract_no = _get_str("合同编号（结账号）")
    project_name = _get_str("项目名称")
    manager = _get_str("项目负责人")
    today = datetime.now().strftime("%Y.%m.%d")

    # --- 1. 处理文档级别段落（报审单编号） ---
    for para in doc.paragraphs:
        if para.text.strip().startswith("报审单编号"):
            # 留空，不填编号
            fill_paragraph(para, "报审单编号：")

    # --- 2. 处理表格内容 ---
    table = doc.tables[0]

    # Row 0: 主内容单元格
    cell = table.rows[0].cells[0]
    for pi, para in enumerate(cell.paragraphs):
        text = para.text
        if not text.strip():
            continue

        # Paragraph 0: 致：...（付款单位）
        if text.strip().startswith("致："):
            new_text = f"致：{payer}" if payer else "致："
            fill_paragraph(para, new_text)

        # Paragraph 1: 根据我单位与贵单位签订的...
        elif "根据我单位与贵单位签订" in text:
            # 智能拼接：避免 "检测" + "检测项目" 重复
            if "检测项目" in project_name:
                name_part = project_name
            elif project_name.endswith("检测"):
                name_part = project_name + "项目"
            elif project_name.endswith("项目") or project_name.endswith("工程"):
                name_part = project_name + "检测"
            else:
                name_part = project_name + "检测项目"
            new_text = (
                f"根据我单位与贵单位签订的{name_part}，"
                f"合同编号为：（{contract_no}），我单位已完成合同约定的全部工作量，"
                f"现报上结算报审单，请贵公司予以审核确认。"
            )
            fill_paragraph(para, new_text)

        # Paragraph 3: 总工作量
        elif "实际发生的总工作量" in text:
            new_text = f"本工程按合同约定单价计费，实际发生的总工作量为{total_workload}元。"
            fill_paragraph(para, new_text)

        # Paragraph 4: 检测费
        elif "实际应支付检测费" in text:
            new_text = f"根据合同的规定，本工程实际应支付检测费为：{test_fee}元。"
            fill_paragraph(para, new_text)

        # Paragraph 5: 已确认收入
        elif "确认收入" in text:
            new_text = f"（其中已确认收入：{confirmed_income}元"
            fill_paragraph(para, new_text)

        # Paragraph 6: 开发票金额 + 已收款金额
        elif "开发票金额" in text:
            new_text = f"开发票金额：{invoice_amount}元  、已收款金额：{received_amount}元）"
            fill_paragraph(para, new_text)

        # Paragraph 7: 项目负责人
        elif "项目负责人" in text:
            new_text = f"项目负责人：{manager}"
            fill_paragraph(para, new_text)

        # Paragraph 19: 日期
        elif "日期：" in text and pi >= 15:
            new_text = f"　　　　　　　　                                 日期：{today}                               "
            fill_paragraph(para, new_text)

    doc.save(output_path)


def main():
    base_dir = os.path.dirname(os.path.abspath(__file__))
    template_path = os.path.join(base_dir, "template.docx")
    excel_path = os.path.join(base_dir, "报审单信息导入模版(1).xlsx")
    output_dir = os.path.join(base_dir, "output")
    os.makedirs(output_dir, exist_ok=True)

    df = pd.read_excel(excel_path, engine="openpyxl")
    print(f"读取到 {len(df)} 条记录")

    for idx, row in df.iterrows():
        contract_no = str(row.get("合同编号（结账号）", "") or "").strip()
        project_name = str(row.get("项目名称", "") or "").strip()
        manager_val = str(row.get("项目负责人", "") or "").strip()

        # 文件名：合同编号+项目名称前20字+项目负责人
        safe_name = project_name[:20].replace("/", "-").replace("\\", "-")
        filename = f"{contract_no}_{safe_name}_{manager_val}.docx"
        output_path = os.path.join(output_dir, filename)

        fill_settlement_form(template_path, output_path, row.to_dict())
        print(f"  [OK] {filename}")

    print(f"\n全部完成，输出目录：{output_dir}")


if __name__ == "__main__":
    main()
