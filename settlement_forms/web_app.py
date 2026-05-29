#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""项目结算报审单 - 网页版

启动: python web_app.py
打开: http://localhost:8080
"""

import os
import sys
import shutil
import zipfile
import tempfile
from datetime import datetime
from pathlib import Path

from fastapi import FastAPI, File, UploadFile, Form
from fastapi.responses import FileResponse, HTMLResponse
from fastapi.staticfiles import StaticFiles
from docx import Document
import pandas as pd

app = FastAPI(title="结算报审单生成器")

# 静态文件
static_dir = Path(__file__).parent / "static"
static_dir.mkdir(exist_ok=True)
app.mount("/static", StaticFiles(directory=static_dir), name="static")

# 内置模板路径（优先使用当前目录下的 template.docx）
BUILTIN_TEMPLATE = Path(__file__).parent / "template.docx"
# 用户上传覆盖的模板
UPLOADED_TEMPLATE = Path(__file__).parent / "uploaded_template.docx"


def get_template_path():
    """获取当前使用的模板路径：优先用户上传的，其次内置的。"""
    if UPLOADED_TEMPLATE.exists():
        return UPLOADED_TEMPLATE
    if BUILTIN_TEMPLATE.exists():
        return BUILTIN_TEMPLATE
    return None


def format_amount(value):
    if pd.isna(value):
        return ""
    v = float(value) * 10000
    if v == int(v):
        return f"{int(v)}"
    return f"{v:.2f}"


def fill_paragraph(paragraph, new_text):
    first_run = None
    for run in paragraph.runs:
        if first_run is None:
            first_run = run
        run.text = ""
    if first_run:
        first_run.text = new_text
    else:
        paragraph.add_run(new_text)


def fill_settlement_form(template_path, output_path, data):
    shutil.copy(template_path, output_path)
    doc = Document(output_path)

    total_workload = format_amount(data.get("实际发生总工作量(含税)", ""))
    test_fee = format_amount(data.get("实际应支付检测费（含税）", ""))
    confirmed_income = format_amount(data.get("已确认收入（含税）", ""))
    invoice_amount = format_amount(data.get("开票金额(含税)", ""))
    received_amount = format_amount(data.get("已收款金额（含税）", ""))

    def _get_str(key):
        val = data.get(key, "")
        if pd.isna(val):
            return ""
        return str(val).strip()

    payer = _get_str("付款单位")
    contract_no = _get_str("合同编号（结账号）")
    project_name = _get_str("项目名称")
    manager = _get_str("项目负责人")
    today = datetime.now().strftime("%Y.%m.%d")

    for para in doc.paragraphs:
        if para.text.strip().startswith("报审单编号"):
            fill_paragraph(para, "报审单编号：")

    table = doc.tables[0]
    cell = table.rows[0].cells[0]
    for pi, para in enumerate(cell.paragraphs):
        text = para.text
        if not text.strip():
            continue

        if text.strip().startswith("致："):
            new_text = f"致：{payer}" if payer else "致："
            fill_paragraph(para, new_text)

        elif "根据我单位与贵单位签订" in text:
            if "检测项目" in project_name:
                name_part = project_name
            elif project_name.endswith("检测"):
                name_part = project_name + "项目"
            elif project_name.endswith("项目") or project_name.endswith("工程"):
                name_part = project_name + "检测"
            else:
                name_part = project_name + "检测项目"
            new_text = (
                f"根据我单位与贵单位签订的{name_part}，"
                f"合同编号为：（{contract_no}），我单位已完成合同约定的全部工作量，"
                f"现报上结算报审单，请贵公司予以审核确认。"
            )
            fill_paragraph(para, new_text)

        elif "实际发生的总工作量" in text:
            new_text = f"本工程按合同约定单价计费，实际发生的总工作量为{total_workload}元。"
            fill_paragraph(para, new_text)

        elif "实际应支付检测费" in text:
            new_text = f"根据合同的规定，本工程实际应支付检测费为：{test_fee}元。"
            fill_paragraph(para, new_text)

        elif "确认收入" in text:
            new_text = f"（其中已确认收入：{confirmed_income}元"
            fill_paragraph(para, new_text)

        elif "开发票金额" in text:
            new_text = f"开发票金额：{invoice_amount}元  、已收款金额：{received_amount}元）"
            fill_paragraph(para, new_text)

        elif "项目负责人" in text:
            new_text = f"项目负责人：{manager}"
            fill_paragraph(para, new_text)

        elif "日期：" in text and pi >= 15:
            new_text = f"　　　　　　　　                                 日期：{today}                               "
            fill_paragraph(para, new_text)

    doc.save(output_path)


@app.get("/", response_class=HTMLResponse)
def index():
    tpl = get_template_path()
    has_builtin = BUILTIN_TEMPLATE.exists()
    has_uploaded = UPLOADED_TEMPLATE.exists()

    if has_uploaded:
        template_status = "已上传（覆盖内置）"
        status_class = "ok"
        tpl_hint = "当前使用上传的模板。留空则使用内置模板。"
    elif has_builtin:
        template_status = "已内置"
        status_class = "ok"
        tpl_hint = "已内置默认模板。如需更换可上传覆盖。"
    else:
        template_status = "未上传"
        status_class = "warn"
        tpl_hint = "请先上传 Word 模板。"

    return f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>项目结算报审单生成器</title>
    <style>
        * {{ box-sizing: border-box; margin: 0; padding: 0; }}
        body {{
            font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, "Helvetica Neue", Arial, sans-serif;
            background: #f5f5f5;
            padding: 40px 20px;
            color: #333;
        }}
        .container {{
            max-width: 600px;
            margin: 0 auto;
            background: #fff;
            border-radius: 12px;
            padding: 40px;
            box-shadow: 0 2px 12px rgba(0,0,0,0.08);
        }}
        h1 {{ font-size: 24px; margin-bottom: 8px; color: #1a1a1a; }}
        .subtitle {{ color: #888; font-size: 14px; margin-bottom: 32px; }}
        .section {{
            margin-bottom: 24px;
            padding: 20px;
            background: #f8f9fa;
            border-radius: 8px;
            border: 1px solid #e8e8e8;
        }}
        .section-title {{
            font-size: 15px;
            font-weight: 600;
            margin-bottom: 12px;
            color: #444;
        }}
        .status {{
            display: inline-block;
            padding: 2px 10px;
            border-radius: 12px;
            font-size: 12px;
            margin-left: 8px;
        }}
        .status.ok {{ background: #d4edda; color: #155724; }}
        .status.warn {{ background: #fff3cd; color: #856404; }}
        input[type="file"] {{
            display: block;
            width: 100%;
            padding: 10px;
            border: 2px dashed #ccc;
            border-radius: 6px;
            background: #fff;
            cursor: pointer;
            transition: border-color 0.2s;
        }}
        input[type="file"]:hover {{ border-color: #999; }}
        button {{
            width: 100%;
            padding: 14px;
            background: #2563eb;
            color: #fff;
            border: none;
            border-radius: 8px;
            font-size: 16px;
            font-weight: 500;
            cursor: pointer;
            transition: background 0.2s;
        }}
        button:hover {{ background: #1d4ed8; }}
        button:disabled {{ background: #ccc; cursor: not-allowed; }}
        .info {{
            font-size: 13px;
            color: #666;
            line-height: 1.6;
            margin-top: 8px;
        }}
        .info code {{
            background: #e8e8e8;
            padding: 1px 5px;
            border-radius: 3px;
            font-size: 12px;
        }}
        #result {{
            margin-top: 20px;
            padding: 16px;
            border-radius: 8px;
            display: none;
        }}
        #result.success {{ display: block; background: #d4edda; color: #155724; }}
        #result.error {{ display: block; background: #f8d7da; color: #721c24; }}
        .loading {{
            display: none;
            text-align: center;
            padding: 20px;
            color: #666;
        }}
        .loading.show {{ display: block; }}
    </style>
</head>
<body>
    <div class="container">
        <h1>项目结算报审单生成器</h1>
        <p class="subtitle">上传 Excel 数据，自动批量生成结算报审单 Word 文档</p>

        <form id="uploadForm" enctype="multipart/form-data">
            <div class="section">
                <div class="section-title">
                    1. Word 模板
                    <span class="status {status_class}">{template_status}</span>
                </div>
                <input type="file" name="template" accept=".docx,.doc">
                <p class="info">{tpl_hint}</p>
            </div>

            <div class="section">
                <div class="section-title">2. Excel 数据文件（必填）</div>
                <input type="file" name="excel" accept=".xlsx,.xls" required>
                <p class="info">Excel 需包含列：<code>合同编号（结账号）</code>、<code>项目名称</code>、<code>实际发生总工作量(含税)</code>、<code>实际应支付检测费（含税）</code>、<code>已确认收入（含税）</code>、<code>开票金额(含税)</code>、<code>已收款金额（含税）</code>、<code>项目负责人</code></p>
            </div>

            <button type="submit" id="submitBtn">生成结算报审单</button>
        </form>

        <div class="loading" id="loading">正在生成，请稍候...</div>
        <div id="result"></div>
    </div>

    <script>
        document.getElementById('uploadForm').addEventListener('submit', async (e) => {{
            e.preventDefault();
            const btn = document.getElementById('submitBtn');
            const loading = document.getElementById('loading');
            const result = document.getElementById('result');

            btn.disabled = true;
            loading.classList.add('show');
            result.className = '';
            result.style.display = 'none';

            const formData = new FormData(e.target);

            try {{
                const res = await fetch('./generate', {{ method: 'POST', body: formData }});
                if (res.ok) {{
                    const blob = await res.blob();
                    const url = window.URL.createObjectURL(blob);
                    const a = document.createElement('a');
                    a.href = url;
                    a.download = '结算报审单.zip';
                    document.body.appendChild(a);
                    a.click();
                    a.remove();
                    window.URL.revokeObjectURL(url);
                    result.className = 'success';
                    result.textContent = '生成成功！文件已开始下载。';
                }} else {{
                    const err = await res.text();
                    result.className = 'error';
                    result.textContent = '生成失败：' + err;
                }}
            }} catch (err) {{
                result.className = 'error';
                result.textContent = '请求出错：' + err.message;
            }} finally {{
                btn.disabled = false;
                loading.classList.remove('show');
            }}
        }});
    </script>
</body>
</html>"""


@app.post("/generate")
def generate(
    excel: UploadFile = File(...),
    template: UploadFile = File(None),
):
    # 保存用户上传的模板（覆盖内置）
    if template and template.filename:
        with open(UPLOADED_TEMPLATE, "wb") as f:
            shutil.copyfileobj(template.file, f)

    tpl_path = get_template_path()
    if not tpl_path:
        return HTMLResponse("未找到模板，请先上传 Word 模板", status_code=400)

    # 读取 Excel
    excel_path = Path(tempfile.gettempdir()) / f"settlement_{datetime.now().strftime('%Y%m%d%H%M%S')}.xlsx"
    with open(excel_path, "wb") as f:
        shutil.copyfileobj(excel.file, f)

    df = pd.read_excel(excel_path, engine="openpyxl")
    if len(df) == 0:
        return HTMLResponse("Excel 文件为空", status_code=400)

    # 生成文件到临时目录
    output_dir = Path(tempfile.gettempdir()) / f"settlement_out_{datetime.now().strftime('%Y%m%d%H%M%S')}"
    output_dir.mkdir(exist_ok=True)

    for idx, row in df.iterrows():
        contract_no = str(row.get("合同编号（结账号）", "") or "").strip()
        project_name = str(row.get("项目名称", "") or "").strip()
        manager_val = str(row.get("项目负责人", "") or "").strip()
        safe_name = project_name[:20].replace("/", "-").replace("\\", "-")
        filename = f"{contract_no}_{safe_name}_{manager_val}.docx"
        output_path = output_dir / filename
        fill_settlement_form(str(tpl_path), str(output_path), row.to_dict())

    # 打包成 zip
    zip_path = Path(tempfile.gettempdir()) / f"settlement_{datetime.now().strftime('%Y%m%d%H%M%S')}.zip"
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for f in output_dir.glob("*.docx"):
            zf.write(f, f.name)

    return FileResponse(
        zip_path,
        media_type="application/zip",
        filename="结算报审单.zip",
    )


if __name__ == "__main__":
    import uvicorn
    port = int(os.environ.get("PORT", 8080))
    uvicorn.run(app, host="0.0.0.0", port=port)
