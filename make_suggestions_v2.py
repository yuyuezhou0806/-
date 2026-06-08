"""生成《项目工作规范》反馈意见 Word文档"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# 全部宋体，标题加粗，正文字号比标题小一号
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)
style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.font.bold = False

for lvl, size in [(1, 15), (2, 14), (3, 13)]:
    hs = doc.styles[f'Heading {lvl}']
    hs.font.name = '宋体'
    hs.font.size = Pt(size)
    hs.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    hs.font.bold = True

title = doc.add_heading('关于《项目工作规范与违规管理条例》的几点反馈建议', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.name = '宋体'
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    run.font.bold = True

doc.add_paragraph()

# 总体
doc.add_heading('一、整体评价', level=1)

doc.add_paragraph(
    '这份规范覆盖了项目全生命周期，从立项到催款闭环完整，权责划分清晰，'
    '对当前项目管理中的常见漏洞（节点滞后、资料缺失、结算卡顿、催款不及时）都做了针对性约束，很有必要。'
)
doc.add_paragraph(
    '提几点具体建议，主要围绕"怎么让一线同事更好落地执行"，供参考。'
)

# 具体建议
doc.add_heading('二、具体建议', level=1)

# 1
doc.add_heading('1. "每日自查"建议改为"异常抽查"', level=2)
doc.add_paragraph(
    '第二十二条要求项目经理和业务人员每日自查，每天至少 30 分钟花在核数据、翻资料上，'
    '对项目多、人手紧的团队来说执行成本很高。建议：正常的项目不查，只在系统数据出现异常'
    '（节点逾期、收入偏差超 5%、资料逾期未传）时推送提醒给对应责任人。把"人找问题"改成"问题找人"，'
    '减少无效重复劳动。'
)

# 2
doc.add_heading('2. 30 天费用录入限制建议给个缓冲', level=2)
doc.add_paragraph(
    '第九条规定劳务费、配合费等超过发生时间 30 天不予审批。这个方向是对的，但现实中常遇到供应商开票慢、'
    '跨月结算单流转延迟等情况，不是项目组故意拖着不录。建议改成阶梯式：30 天提醒、60 天预警需说明原因、'
    '90 天拦截不予审批。中间给一个合理的缓冲窗口。'
)

# 3
doc.add_heading('3. 电子+纸质双归档建议分层执行', level=2)
doc.add_paragraph(
    '第十条要求所有项目电子+纸质双归档。对内部项目而言每份资料都要打印、装订、理档，'
    '工作量不小。建议：甲方合同明确要求纸质归档的才走双轨，其他项目纯电子归档即可。'
    '既减轻一线负担，也节省公司档案室空间。'
)

# 4
doc.add_heading('4. 处罚梯度中建议加入"整改期"', level=2)
doc.add_paragraph(
    '第二十三条从一般违规到重大失职的处罚梯度很清楚。建议在每档处罚前加一个"整改期"——'
    '比如中度违规先给予 3-5 个工作日整改窗口，整改到位的从轻处理或免于追责。'
    '不给人一次改正机会直接处罚，可能导致大家怕担责反而隐瞒问题。'
)

# 5
doc.add_heading('5. 结算流转环节建议减少人工交接', level=2)
doc.add_paragraph(
    '第十一条结算单的传递链很长——项目经理编→子公司检→系统发起→用印→通知业务员→业务员领取→送达甲方，'
    '每多一个交接点就多一个卡顿风险。建议用印后的"通知→领取"两步用系统自动推送替代人工通知。'
    '业务员手机收到推送直接去领就行，不会漏。'
)

# 6
doc.add_heading('6. 催款模块建议加自动提醒', level=2)
doc.add_paragraph(
    '第十二条催收要求业务员主动跟进，但如果一个人管十几个项目，付款节点一多容易漏。'
    '建议：系统根据合同付款节点到期前自动弹窗提醒，逾期自动抄送主管。'
    '业务员不用自己记时间，系统推着走。'
)

# 7
doc.add_heading('7. 资料上传建议做前端校验', level=2)
doc.add_paragraph(
    '第十条对资料完整性要求很细（签字、盖章、日期、影像佐证），但如果上传之后再退回来补就很耽误时间。'
    '建议在提交界面做简单的校验提示——比如缺盖章提醒、日期缺失弹窗。'
    '在源头卡住残缺资料，省得返工。'
)

# 总结
doc.add_heading('三、总结', level=1)
doc.add_paragraph(
    '以上 7 条建议的核心思路是同一个：规范要严，但执行要顺。'
    '把重心从"事后处罚"往前挪到"事前预防"和"事中提醒"，'
    '让系统辅助人而不是人伺候系统，规范才能真正落地。'
)

doc.add_paragraph()
p_end = doc.add_paragraph()
p_end.add_run('以上建议供讨论稿参考。\n\n2026 年 6 月 2 日')

output = r'C:\Users\admin\Desktop\项目规范反馈建议.docx'
doc.save(output)
print(f'Done: {output}')
